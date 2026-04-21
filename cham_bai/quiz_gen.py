from __future__ import annotations

import json
import random
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from cham_bai.openrouter import ChatMessage, complete_chat
from cham_bai.quiz_excel import (
    QUIZ_FIXED_LEVELS,
    VerticalQuizBlock,
    _norm_key,
    clean_quiz_explanation,
    fill_template_from_rows,
    fill_template_session_warmup_quiz,
    fill_template_vertical_quiz,
    is_vertical_quiz_template,
    read_headers_from_template,
)
from cham_bai.session_warmup_plan import apply_session_warmup_plan
from cham_bai.session_end_plan import apply_session_end_plan
from cham_bai.settings import model as resolve_model


def sanitize_quiz_filename_part(s: str, max_len: int = 48) -> str:
    """Tên file thân thiện Windows: bỏ ký tự cấm, rút gọn."""
    s = (s or "").strip()
    if not s:
        return "x"
    s = re.sub(r"[\x00-\x1f]", "", s)
    s = re.sub(r'[<>:"/\\|?*]', "_", s)
    s = re.sub(r"\s+", "_", s)
    s = s.strip("._") or "x"
    return s[:max_len]


def default_quiz_output_path(template_xlsx: Path, lesson: str, session: str) -> Path:
    """
    Đường dẫn .xlsx mới trong cùng thư mục với file mẫu:
    quiz_<lesson>_<session>_<YYYYMMDD_HHMMSS>.xlsx (tránh trùng tên).
    """
    base = template_xlsx.resolve().parent
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    sl = sanitize_quiz_filename_part(lesson)
    se = sanitize_quiz_filename_part(session)
    stem = f"quiz_{sl}_{se}_{ts}"
    p = base / f"{stem}.xlsx"
    n = 1
    while p.exists():
        p = base / f"{stem}_{n}.xlsx"
        n += 1
    return p


def default_session_warmup_quiz_output_path(template_xlsx: Path, session_current: str) -> Path:
    """
    Đường dẫn .xlsx quiz session đầu giờ (cùng thư mục với mẫu):
    quizz_session_Dau_gio_<Tên session hiện tại>_<YYYYMMDD_HHMMSS>.xlsx (tránh trùng tên).
    """
    base = template_xlsx.resolve().parent
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    part = sanitize_quiz_filename_part(session_current)
    stem = f"quizz_session_Dau_gio_{part}_{ts}"
    p = base / f"{stem}.xlsx"
    n = 1
    while p.exists():
        p = base / f"{stem}_{n}.xlsx"
        n += 1
    return p


def default_session_end_quiz_output_path(template_xlsx: Path, session_current: str) -> Path:
    """
    Đường dẫn .xlsx quiz session cuối giờ (cùng thư mục với mẫu):
    quizz_session_Cuoi_gio_<Tên session hiện tại>_<YYYYMMDD_HHMMSS>.xlsx (tránh trùng tên).
    """
    base = template_xlsx.resolve().parent
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    part = sanitize_quiz_filename_part(session_current)
    stem = f"quizz_session_Cuoi_gio_{part}_{ts}"
    p = base / f"{stem}.xlsx"
    n = 1
    while p.exists():
        p = base / f"{stem}_{n}.xlsx"
        n += 1
    return p


def _extract_docx_plain(path: str | Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _strip_markdown_fence(text: str) -> str:
    raw = text.strip()
    if "```" not in raw:
        return raw
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", raw, re.I)
    if m:
        return m.group(1).strip()
    return raw


def _extract_first_json_array(raw: str) -> str:
    """Trích mảng JSON ngoài cùng, tôn trọng chuỗi (dấu [ ] trong chuỗi không làm sai lệch)."""
    start = raw.find("[")
    if start == -1:
        raise ValueError("Không tìm thấy ký tự '[' mở mảng JSON.")
    depth = 0
    in_string = False
    escape = False
    i = start
    while i < len(raw):
        c = raw[i]
        if in_string:
            if escape:
                escape = False
            elif c == "\\":
                escape = True
            elif c == '"':
                in_string = False
            i += 1
            continue
        if c == '"':
            in_string = True
            i += 1
            continue
        if c == "[":
            depth += 1
        elif c == "]":
            depth -= 1
            if depth == 0:
                return raw[start : i + 1]
        i += 1
    raise ValueError("Mảng JSON chưa đóng hoặc không hợp lệ.")


def _parse_json_array(text: str) -> list[Any]:
    raw = _strip_markdown_fence(text).strip()
    err: Exception | None = None
    try:
        snippet = _extract_first_json_array(raw)
        data = json.loads(snippet)
        if isinstance(data, list):
            return data
    except Exception as e:
        err = e
    try:
        obj = json.loads(raw)
        if isinstance(obj, list):
            return obj
        for k in ("questions", "items", "data", "quiz"):
            v = obj.get(k) if isinstance(obj, dict) else None
            if isinstance(v, list):
                return v
    except Exception as e:
        if err is None:
            err = e
    raise ValueError(str(err) if err else "Không parse được JSON.")


def _finish_reason(data: dict[str, Any]) -> str | None:
    try:
        v = data["choices"][0].get("finish_reason")
        return v if isinstance(v, str) else None
    except (KeyError, IndexError, TypeError):
        return None


QUIZ_KIND_SESSION = "session"
QUIZ_KIND_LESSON = "lesson"
QUIZ_KIND_SESSION_WARMUP = "session_warmup"
QUIZ_KIND_SESSION_END = "session_end"


def normalize_quiz_kind(kind: str | None) -> str:
    k = (kind or QUIZ_KIND_SESSION).strip().lower()
    if k == QUIZ_KIND_LESSON:
        return QUIZ_KIND_LESSON
    if k == QUIZ_KIND_SESSION_WARMUP:
        return QUIZ_KIND_SESSION_WARMUP
    if k == QUIZ_KIND_SESSION_END:
        return QUIZ_KIND_SESSION_END
    return QUIZ_KIND_SESSION


def _quiz_kind_bullet(kind: str) -> str:
    if kind == QUIZ_KIND_LESSON:
        return (
            "- Loại: QUIZ THEO LESSON — câu hỏi tổng hợp/ôn cả lesson; "
            "bám tên lesson và DOCX (nếu có), có thể rộng hơn phạm vi một session cụ thể.\n"
        )
    if kind == QUIZ_KIND_SESSION_WARMUP:
        return (
            "- Loại: QUIZZ SESSION ĐẦU GIỜ — kiểm tra nhanh trước khi vào bài; "
            "ưu tiên câu hỏi ngắn, sát \"session hiện tại\", có 1–2 câu nối kiến thức từ \"session trước\".\n"
        )
    if kind == QUIZ_KIND_SESSION_END:
        return (
            "- Loại: QUIZZ SESSION CUỐI GIỜ — tổng kết/đánh giá nhanh theo session hiện tại; "
            "ưu tiên câu hỏi sát nội dung vừa học.\n"
        )
    return (
        "- Loại: QUIZ THEO SESSION — quiz đầu giờ sát nội dung session, "
        "bám cặp lesson + session và DOCX (nếu có).\n"
    )


def _subject_bullet(subject: str) -> str:
    s = (subject or "").strip()
    if not s:
        return ""
    return (
        f"- Môn học / lĩnh vực: «{s}» — toàn bộ câu hỏi, ví dụ, code và thuật ngữ phải **nhất quán** với môn này "
        f"(ngôn ngữ/công nghệ đúng domain, độ khó phù hợp đại học).\n"
    )


SYSTEM_QUIZ = """Bạn là giảng viên đại học. Soạn câu hỏi trắc nghiệm quiz — nội dung tiếng Việt.

Chỉ trả về một mảng JSON hợp lệ. Không markdown, không ```, không chữ ngoài mảng.

Mỗi phần tử là object với ĐÚNG các khóa sau (chỉ ASCII, không dùng tên cột Excel tiếng Việt làm khóa — tránh trùng khóa JSON):
- "stt": số nguyên
- "q": chuỗi — nội dung câu hỏi
- "a", "b", "c", "d": chuỗi — bốn đáp án (bốn lựa chọn phải khác nhau, mạch lạc)
- "ans": một ký tự "A", "B", "C" hoặc "D" — đáp án đúng

Nếu trong "q" có đoạn code: **tên biến/hàm/identifier chỉ tiếng Anh** (snake_case); không dùng tên romanized tiếng Việt trong code.

Nếu user cung cấp môn học: bám đúng lĩnh vực đó (không lệch sang môn khác).

Đủ N phần tử, stt từ 1 đến N. Chương trình có thể xáo thứ tự A–D khi ghi file để phân bố đáp án đúng ngẫu nhiên."""


SYSTEM_QUIZ_VERTICAL = """Bạn là giảng viên đại học. Nội dung câu hỏi và đáp án bằng tiếng Việt.

QUAN TRỌNG — Định dạng JSON:
- Chỉ trả về MỘT mảng JSON: bắt đầu bằng [ và KẾT THÚC bằng ] trên cùng một lần trả lời.
- TUYỆT ĐỐI KHÔNG dùng khóa tiếng Việt (không dùng "Câu hỏi", "Mức độ", "Giải thích" làm tên khóa). CHỈ được dùng đúng 5 khóa ASCII: q, options, t, ok, e.
- Không markdown, không ```, không văn bản trước/sau mảng.

Đúng 5 phần tử (câu 1→5). Mỗi phần tử:
- "q": chuỗi — nội dung câu hỏi (code được dùng \\n)
- "options": đúng 4 object (thứ tự A→D), mỗi object CHỈ có "t", "ok", "e":
  - "t": nội dung đáp án
  - "ok": true hoặc false — đúng MỘT trong bốn phải là true
  - "e": bắt buộc, không rỗng — CHỈ viết lý do (vì sao đúng / vì sao sai). KHÔNG ghi tiền tố "Đúng:", "Sai:", "Đúng —", "Sai —" (cột Kết quả trên Excel đã có Đúng/Sai riêng).
  - Giữ mỗi "e" gọn (khoảng 80–140 ký tự) để đủ chỗ kết thúc JSON.

Nếu trong "q" có đoạn code Python minh họa: logic phải đúng, khớp đáp án; **tên biến/hàm/identifier trong code chỉ tiếng Anh** (vd. `status`, `order_id`), không romanized Việt (`trang_thai`, `ma`, …). Tránh lỗi kiểu dùng continue hoặc if khiến thiếu bước/ thiếu giá trị (ví dụ: for i in range(1,6): if i%2==0: continue; print(i) chỉ in 1,3,5 — không hỏi kiểu "in ra 1,3,4" vì sẽ không khớp code). Code mẫu phải chạy đúng và thống nhất với các lựa chọn.

Nếu user cung cấp môn học: mọi câu hỏi và ví dụ phải phù hợp môn/lĩnh vực đó (không lệch domain).

Thứ tự 4 object trong "options" do bạn soạn; chương trình sẽ xáo lại thứ tự A–D khi ghi Excel để đáp án đúng rơi ngẫu nhiên vào A/B/C/D (tránh cả bài trùng một vị trí đúng).

Tên Mức độ trên file Excel do chương trình điền đúng 5 mức: «Thông hiểu», «Thông hiểu», «Vận dụng sơ bộ», «Phân biệt/So sánh», «Phân tích sơ bộ» — không dùng nhãn như Nhớ, Hiểu, Vận dụng."""


SYSTEM_QUIZ_SESSION_WARMUP = (
    "Bạn là giảng viên đại học. Soạn quiz đầu giờ (kiểm tra bài cũ + chuẩn bị bài mới) — nội dung tiếng Việt.\n\n"
    "Chỉ trả về một mảng JSON hợp lệ. Không markdown, không ```, không chữ ngoài mảng.\n\n"
    "Yêu cầu đầu ra:\n"
    "- Đúng 45 phần tử.\n"
    "- 30 câu đầu: nội dung sát session trước (BÀI CŨ); 15 câu sau: sát session hiện tại (BÀI MỚI).\n"
    "- Mỗi phần tử là object với ĐÚNG các khóa ASCII: part, question_content, answers (4 string), explanations (4 string), isCorrect (1..4), difficulty (số nguyên).\n"
    "- \"part\": \"prev\" cho 30 câu đầu, \"current\" cho 15 câu sau.\n"
    "- \"difficulty\": CHỈ được dùng các số: 4, 5, 6, 7, 8, 9.\n"
    "  (4 Vận dụng chuyên sâu; 5 Phân tích chuyên sâu; 6 Sáng tạo; 7 Thông hiểu; 8 Vận dụng sơ bộ; 9 Phân tích sơ bộ)\n"
    "- Cột category trên Excel do tool điền (BÀI CŨ / BÀI MỚI); không thêm khóa category trong JSON.\n\n"
    "Ràng buộc:\n"
    "- Giải thích phải đủ rõ vì sao đúng/sai (ngắn gọn, không rỗng).\n"
    "- Nếu có code trong câu hỏi: identifier chỉ tiếng Anh (snake_case), logic khớp đáp án.\n"
    "- Mục tiêu: kiểm tra bài cũ + sinh viên đã chuẩn bị bài mới, đúng phân bổ từng vị trí."
)

SYSTEM_QUIZ_SESSION_END = (
    "Bạn là giảng viên đại học. Soạn quiz cuối giờ (tổng kết/đánh giá nhanh) — nội dung tiếng Việt.\n\n"
    "Chỉ trả về một mảng JSON hợp lệ. Không markdown, không ```, không chữ ngoài mảng.\n\n"
    "Yêu cầu đầu ra:\n"
    "- Đúng 45 phần tử.\n"
    "- TẤT CẢ câu hỏi đều bám sát session hiện tại (BÀI MỚI).\n"
    "- Mỗi phần tử là object với ĐÚNG các khóa ASCII: part, question_content, answers (4 string), explanations (4 string), isCorrect (1..4), difficulty (số nguyên).\n"
    "- \"part\": luôn là \"current\".\n"
    "- \"difficulty\": CHỈ được dùng các số: 6, 10, 11.\n"
    "  (6 Sáng tạo; 10 Vận dụng; 11 Phân tích) — tool sẽ tự chia đều.\n\n"
    "Ràng buộc:\n"
    "- Giải thích đủ rõ vì sao đúng/sai (ngắn gọn, không rỗng).\n"
    "- Nếu có code: identifier chỉ tiếng Anh (snake_case), logic khớp đáp án.\n"
    "- Nếu code có thụt lề (Python/SQL): PHẢI giữ đúng thụt lề bằng 4 dấu cách. Ví dụ:\n"
    "  Code:\\n"
    "  for i in range(3):\\n"
    "      print(i)\\n"
    "  (không được mất khoảng trắng đầu dòng như 'print(i)' nằm sát lề).\n"
    "- Mục tiêu: bám sát session hiện tại, không lẫn session trước."
)


def _end_retry_messages(bad_raw: str) -> list[ChatMessage]:
    tail = bad_raw[-9000:] if len(bad_raw) > 9000 else bad_raw
    return [
        ChatMessage(
            role="system",
            content=(
                "Chỉ output DUY NHẤT một mảng JSON hợp lệ gồm đúng 45 object. "
                "Không markdown, không ```, không chữ ngoài mảng. "
                "Mỗi object có đúng các khóa ASCII: part, question_content, answers, explanations, isCorrect, difficulty. "
                "answers và explanations là mảng đúng 4 string; isCorrect là 1..4; "
                "difficulty chỉ được là 6/10/11; part luôn là 'current'. "
                "Toàn bộ output phải bắt đầu bằng [ và kết thúc bằng ]."
            ),
        ),
        ChatMessage(
            role="user",
            content=(
                "Output trước bị thiếu/cụt hoặc sai định dạng. "
                "Hãy output LẠI TOÀN BỘ mảng 45 phần tử từ dấu [ đến dấu ] (đừng cố nối tiếp).\n\n"
                f"Output trước (có thể bị cắt):\n{tail}"
            ),
        ),
    ]


def _end_block_messages(
    *,
    subject: str,
    session_current: str,
    start_stt: int,
    n: int,
    docx_excerpt: str,
) -> list[ChatMessage]:
    end_stt = start_stt + n - 1
    user = (
        "Thông tin:\n"
        f"{_subject_bullet(subject)}"
        f"- Session hiện tại: {session_current}\n\n"
        f"Nhiệm vụ: soạn đúng {n} câu, STT {start_stt}–{end_stt} (KHÔNG ghi STT vào JSON), "
        "tất cả đều thuộc session hiện tại.\n"
        "Yêu cầu: mọi object phải có part='current'.\n"
    )
    if docx_excerpt.strip():
        user += f"\nNội dung bài giảng (trích từ tài liệu cung cấp):\n---\n{docx_excerpt}\n---\n"
        user += (
            "\nRàng buộc nội dung: chỉ tạo câu hỏi dựa trên thông tin có trong tài liệu. "
            "Nếu tài liệu không có thì KHÔNG được tự bịa thêm ngoài chương trình.\n"
        )
    sys = (
        "Chỉ output DUY NHẤT một mảng JSON hợp lệ gồm đúng "
        f"{n} object. Không markdown, không ```, không chữ ngoài mảng.\n"
        "Mỗi object có đúng các khóa ASCII: part, question_content, answers, explanations, isCorrect, difficulty.\n"
        "answers và explanations là mảng đúng 4 string; isCorrect là 1..4; difficulty chỉ được là 6/10/11; part luôn là 'current'.\n"
        "Để tránh bị cắt output: viết NGẮN — question_content <= 220 ký tự; mỗi explanation <= 120 ký tự.\n"
        "Nếu có code: đặt code ở CUỐI question_content theo đúng cấu trúc:\n"
        "Code:\\n<dòng 1>\\n<dòng 2>... (giữ thụt lề chuẩn bằng 4 dấu cách, không dùng markdown fence). "
        "Chỉ question_content được phép có xuống dòng; answers/explanations phải 1 dòng.\n"
        "Output phải bắt đầu bằng [ và kết thúc bằng ]."
    )
    return [ChatMessage(role="system", content=sys), ChatMessage(role="user", content=user)]


def _end_block_retry_messages(*, bad_raw: str, n: int) -> list[ChatMessage]:
    tail = bad_raw[-9000:] if len(bad_raw) > 9000 else bad_raw
    return [
        ChatMessage(
            role="system",
            content=(
                "Chỉ output DUY NHẤT một mảng JSON hợp lệ gồm đúng "
                f"{n} object. Không markdown, không ```, không chữ ngoài mảng. "
                "Mỗi object có đúng các khóa ASCII: part, question_content, answers, explanations, isCorrect, difficulty. "
                "answers và explanations là mảng đúng 4 string; isCorrect là 1..4; difficulty chỉ được là 6/10/11; part luôn là 'current'. "
                "Output phải bắt đầu bằng [ và kết thúc bằng ]."
            ),
        ),
        ChatMessage(
            role="user",
            content=(
                "Output trước bị thiếu/cụt hoặc sai định dạng. "
                "Hãy output LẠI TOÀN BỘ mảng (đừng cố nối tiếp).\n\n"
                f"Output trước (có thể bị cắt):\n{tail}"
            ),
        ),
    ]


def _vertical_retry_messages(bad_raw: str) -> list[ChatMessage]:
    tail = bad_raw[-7500:] if len(bad_raw) > 7500 else bad_raw
    return [
        ChatMessage(
            role="system",
            content=(
                "Output ONLY a complete valid JSON array of exactly 5 objects. "
                "Each object must be: "
                '{"q":"...","options":['
                '{"t":"...","ok":false,"e":"..."},'
                '{"t":"...","ok":true,"e":"..."},'
                '{"t":"...","ok":false,"e":"..."},'
                '{"t":"...","ok":false,"e":"..."}'
                "]} "
                "Use ONLY these ASCII keys: q, options, t, ok, e. "
                "Exactly 4 options per question; exactly one ok:true per question. "
                "Each option MUST have non-empty \"e\" with reasoning only — NO prefix \"Correct:\" or \"Wrong:\". "
                "Each e under 120 characters. No markdown, no code fences. "
                "Your entire response must start with [ and end with ]."
            ),
        ),
        ChatMessage(
            role="user",
            content=(
                "The previous output was invalid or truncated. Output the FULL array again "
                "from the opening [ through the closing ].\n\n"
                f"Previous output (possibly incomplete):\n{tail}"
            ),
        ),
    ]


def _flat_retry_messages(bad_raw: str, n: int) -> list[ChatMessage]:
    tail = bad_raw[-6000:] if len(bad_raw) > 6000 else bad_raw
    return [
        ChatMessage(
            role="system",
            content=(
                f"Output ONLY a valid JSON array of exactly {n} objects. "
                'Each: {"stt":int,"q":"...","a":"...","b":"...","c":"...","d":"...","ans":"A"|"B"|"C"|"D"}. '
                "Keys only: stt, q, a, b, c, d, ans. No markdown. Start with [ end with ]."
            ),
        ),
        ChatMessage(
            role="user",
            content=(
                "Previous output was invalid or cut off. Output the complete array again.\n\n"
                f"{tail}"
            ),
        ),
    ]


def _warmup_retry_messages(bad_raw: str) -> list[ChatMessage]:
    tail = bad_raw[-9000:] if len(bad_raw) > 9000 else bad_raw
    return [
        ChatMessage(
            role="system",
            content=(
                "Chỉ output DUY NHẤT một mảng JSON hợp lệ gồm đúng 45 object. "
                "Không markdown, không ```, không chữ ngoài mảng. "
                "Mỗi object có đúng các khóa ASCII: part, question_content, answers, explanations, isCorrect, difficulty. "
                "answers và explanations là mảng đúng 4 string; isCorrect là 1..4; difficulty chỉ được là 4/5/6/7/8/9. "
                "30 câu đầu part='prev', 15 câu sau part='current'. "
                "Toàn bộ output phải bắt đầu bằng [ và kết thúc bằng ]."
            ),
        ),
        ChatMessage(
            role="user",
            content=(
                "Output trước bị thiếu/cụt hoặc sai định dạng. "
                "Hãy output LẠI TOÀN BỘ mảng 45 phần tử từ dấu [ đến dấu ] (đừng cố nối tiếp). "
                "Giữ đúng thứ tự 45 câu theo bảng vị trí đã đưa.\n\n"
                f"Output trước (có thể bị cắt):\n{tail}"
            ),
        ),
    ]


def _warmup_block_messages(
    *,
    subject: str,
    session_prev: str,
    session_current: str,
    part: str,
    start_stt: int,
    n: int,
    docx_excerpt: str,
) -> list[ChatMessage]:
    part_norm = (part or "").strip().lower()
    if part_norm not in ("prev", "current"):
        part_norm = "prev"
    end_stt = start_stt + n - 1
    focus = (
        f"session trước: «{session_prev}»" if part_norm == "prev" else f"session hiện tại: «{session_current}»"
    )
    user = (
        "Thông tin:\n"
        f"{_subject_bullet(subject)}"
        f"- Session trước: {session_prev}\n"
        f"- Session hiện tại: {session_current}\n\n"
        f"Nhiệm vụ: soạn đúng {n} câu, STT {start_stt}–{end_stt} (KHÔNG ghi STT vào JSON), thuộc {focus}.\n"
        f"Yêu cầu: mọi object phải có part='{part_norm}'.\n"
    )
    if docx_excerpt.strip():
        user += f"\nNội dung bài giảng (trích từ tài liệu cung cấp):\n---\n{docx_excerpt}\n---\n"
        user += (
            "\nRàng buộc nội dung: chỉ tạo câu hỏi dựa trên thông tin có trong tài liệu. "
            "Nếu tài liệu không có thì KHÔNG được tự bịa thêm ngoài chương trình.\n"
        )
    sys = (
        "Chỉ output DUY NHẤT một mảng JSON hợp lệ gồm đúng "
        f"{n} object. Không markdown, không ```, không chữ ngoài mảng.\n"
        "Mỗi object có đúng các khóa ASCII: part, question_content, answers, explanations, isCorrect, difficulty.\n"
        "answers và explanations là mảng đúng 4 string; isCorrect là 1..4; difficulty chỉ được là 4/5/6/7/8/9.\n"
        "Để tránh bị cắt output: viết NGẮN — question_content <= 220 ký tự; "
        "mỗi explanation <= 120 ký tự.\n"
        "Nếu có code: đặt code ở CUỐI question_content theo đúng cấu trúc:\n"
        "Code:\\n<dòng 1>\\n<dòng 2>... (giữ thụt lề chuẩn, không dùng markdown fence). "
        "Chỉ question_content được phép có xuống dòng; answers/explanations phải 1 dòng.\n"
        "Output phải bắt đầu bằng [ và kết thúc bằng ]."
    )
    return [ChatMessage(role="system", content=sys), ChatMessage(role="user", content=user)]


def _warmup_block_retry_messages(*, bad_raw: str, n: int, part: str) -> list[ChatMessage]:
    tail = bad_raw[-9000:] if len(bad_raw) > 9000 else bad_raw
    part_norm = (part or "").strip().lower()
    if part_norm not in ("prev", "current"):
        part_norm = "prev"
    return [
        ChatMessage(
            role="system",
            content=(
                "Chỉ output DUY NHẤT một mảng JSON hợp lệ gồm đúng "
                f"{n} object. Không markdown, không ```, không chữ ngoài mảng. "
                "Mỗi object có đúng các khóa ASCII: part, question_content, answers, explanations, isCorrect, difficulty. "
                "answers và explanations là mảng đúng 4 string; isCorrect là 1..4; difficulty là số nguyên 4..11. "
                "Viết NGẮN để không bị cắt: question_content <= 220 ký tự; mỗi explanation <= 120 ký tự. "
                "Nếu có code: đặt ở cuối question_content theo dạng 'Code:\\n...' (không markdown fence). "
                "Chỉ question_content được phép có xuống dòng; answers/explanations phải 1 dòng. "
                f"Mọi object phải có part='{part_norm}'. "
                "Output phải bắt đầu bằng [ và kết thúc bằng ]."
            ),
        ),
        ChatMessage(
            role="user",
            content=(
                "Output trước bị thiếu/cụt hoặc sai định dạng. "
                "Hãy output LẠI TOÀN BỘ mảng (đừng cố nối tiếp). "
                f"Phải đúng {n} phần tử.\n\n"
                f"Output trước (có thể bị cắt):\n{tail}"
            ),
        ),
    ]


def _parse_session_warmup_items(arr: list[Any], *, plan: str = "warmup") -> list[dict[str, object]]:
    if len(arr) < 45:
        raise ValueError(f"Cần đúng 45 câu, model trả {len(arr)} phần tử.")
    if len(arr) > 45:
        # Model đôi khi trả thừa; lấy 45 phần tử đầu (giữ thứ tự khớp bảng vị trí).
        arr = arr[:45]
    out: list[dict[str, object]] = []
    for i, it in enumerate(arr):
        if not isinstance(it, dict):
            raise ValueError(f"Câu {i + 1}: mỗi phần tử phải là object.")
        part = str(it.get("part", "")).strip().lower()
        if plan == "warmup":
            if i < 30 and part not in ("prev", ""):
                raise ValueError(f"Câu {i + 1}: 30 câu đầu phải part='prev'.")
            if i >= 30 and part not in ("current", ""):
                raise ValueError(f"Câu {i + 1}: 15 câu sau phải part='current'.")
        elif plan == "end":
            if part not in ("current", ""):
                raise ValueError(f"Câu {i + 1}: part phải là 'current'.")
        else:
            raise ValueError("Plan không hợp lệ.")
        q = it.get("question_content")
        if not isinstance(q, str) or not q.strip():
            raise ValueError(f"Câu {i + 1}: thiếu question_content.")
        answers = it.get("answers")
        exps = it.get("explanations")
        if not isinstance(answers, list) or len(answers) != 4 or not all(isinstance(x, str) and x.strip() for x in answers):
            raise ValueError(f"Câu {i + 1}: answers phải là mảng 4 chuỗi.")
        if not isinstance(exps, list) or len(exps) != 4 or not all(isinstance(x, str) and x.strip() for x in exps):
            raise ValueError(f"Câu {i + 1}: explanations phải là mảng 4 chuỗi.")
        is_correct = it.get("isCorrect")
        try:
            ic = int(is_correct)
        except Exception:
            raise ValueError(f"Câu {i + 1}: isCorrect phải là số 1..4.")
        if ic not in (1, 2, 3, 4):
            raise ValueError(f"Câu {i + 1}: isCorrect phải là số 1..4.")
        row: dict[str, object] = {
            "question_content": q.strip(),
            "answer_1": str(answers[0]).strip(),
            "explanation_answer_1": str(exps[0]).strip(),
            "answer_2": str(answers[1]).strip(),
            "explanation_answer_2": str(exps[1]).strip(),
            "answer_3": str(answers[2]).strip(),
            "explanation_answer_3": str(exps[2]).strip(),
            "answer_4": str(answers[3]).strip(),
            "explanation_answer_4": str(exps[3]).strip(),
            "isCorrect": ic,
            "difficulty": 4,
        }
        out.append(row)
    if plan == "warmup":
        apply_session_warmup_plan(out)
    else:
        apply_session_end_plan(out)
    return out


def _flex_question_text(item: dict[str, Any]) -> str | None:
    for k in ("q", "question", "question_text", "cau_hoi", "text"):
        v = item.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return None


def _flex_options_list(item: dict[str, Any]) -> list[Any] | None:
    v = item.get("options")
    if isinstance(v, list) and v:
        return v
    v = item.get("choices")
    if isinstance(v, list) and v:
        return v
    return None


def _flex_opt_text(o: dict[str, Any]) -> str | None:
    for k in ("t", "text", "choice", "noi_dung", "label"):
        v = o.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return None


def _flex_opt_explain(o: dict[str, Any]) -> str:
    for k in ("e", "explain", "giai_thich", "why"):
        v = o.get(k)
        if isinstance(v, str):
            return v.strip()
    return ""


def _flex_opt_ok(o: dict[str, Any]) -> bool | None:
    if "ok" in o:
        v = o["ok"]
        if isinstance(v, bool):
            return v
    if "correct" in o:
        v = o["correct"]
        if isinstance(v, bool):
            return v
        if v in (1, "1", "true", "True", "yes"):
            return True
        if v in (0, "0", "false", "False", "no"):
            return False
    if "dung" in o:
        v = o["dung"]
        if isinstance(v, bool):
            return v
    return None


def _pick_header(headers: list[str], *candidates: str) -> str | None:
    norms = {_norm_key(h): h for h in headers}
    for c in candidates:
        k = _norm_key(c)
        if k in norms:
            return norms[k]
    return None


def flat_english_items_to_rows(items: list[dict[str, Any]], headers: list[str]) -> list[dict[str, Any]]:
    """Ánh xạ object JSON (stt, q, a–d, ans) sang dict theo đúng tên cột file mẫu."""
    h_stt = _pick_header(headers, "STT", "stt", "Stt")
    h_q = _pick_header(headers, "Câu hỏi", "Cau hoi")
    letters = []
    for lab in ("A", "B", "C", "D"):
        hl = _pick_header(
            headers,
            f"Đáp án {lab}",
            f"Đáp án{lab}",
            f"Dap an {lab}",
            f"Đáp án {lab.lower()}",
            f"dap an {lab.lower()}",
        )
        letters.append(hl)
    h_ans = _pick_header(
        headers,
        "Đáp án đúng",
        "Dap an dung",
        "Đápánđúng",
    )
    rows_out: list[dict[str, Any]] = []
    for it in items:
        row: dict[str, Any] = {}
        if h_stt is not None:
            row[h_stt] = it.get("stt", it.get("n"))
        if h_q is not None:
            row[h_q] = it.get("q") or it.get("question")
        for j, lab in enumerate(("a", "b", "c", "d")):
            if letters[j] is not None:
                row[letters[j]] = it.get(lab) or it.get(lab.upper())
        if h_ans is not None:
            ans = it.get("ans") or it.get("answer") or it.get("dung")
            if ans is not None:
                row[h_ans] = str(ans).strip().upper()[:1]
        rows_out.append(row)
    return rows_out


@dataclass
class QuizGenParams:
    template_xlsx: Path
    docx_path: Path | None
    lesson: str
    session: str
    num_questions: int
    model: str
    output_xlsx: Path
    # Tài liệu bài giảng (ưu tiên Google Docs text từ web). Nếu rỗng thì sẽ fallback docx_path.
    lecture_text: str = ""
    subject: str = ""
    session_prev: str = ""
    session_current: str = ""
    temperature: float = 0.35
    max_tokens: int = 8192
    quiz_kind: str = QUIZ_KIND_SESSION


def _random_correct_slots_five() -> list[int]:
    """5 vị trí hàng đáp án đúng (0=A … 3=D): gồm đủ 0..3 + một vị trí ngẫu nhiên, xáo để phân bố đều hơn."""
    slots = [0, 1, 2, 3, random.randint(0, 3)]
    random.shuffle(slots)
    return slots


def _reorder_vertical_choices_to_slot(choices: tuple[tuple[str, str, str], ...], target_row: int) -> tuple[
    tuple[str, str, str],
    tuple[str, str, str],
    tuple[str, str, str],
    tuple[str, str, str],
]:
    """Đặt đáp án Đúng vào hàng target_row (0–3), xáo ngẫu nhiên 3 đáp án Sai ở các hàng còn lại."""
    rows = list(choices)
    correct_i = next(i for i, r in enumerate(rows) if r[1] == "Đúng")
    correct = rows[correct_i]
    wrong = [r for i, r in enumerate(rows) if i != correct_i]
    random.shuffle(wrong)
    out: list[tuple[str, str, str] | None] = [None, None, None, None]
    out[target_row] = correct
    wi = 0
    for pos in range(4):
        if pos == target_row:
            continue
        out[pos] = wrong[wi]
        wi += 1
    if not all(x is not None for x in out):
        raise RuntimeError("shuffle vertical choices: internal error")
    return (out[0], out[1], out[2], out[3])  # type: ignore[return-value]


def _shuffle_flat_item_answers(it: dict[str, Any]) -> dict[str, Any]:
    """Xáo thứ tự a–d, cập nhật ans tương ứng (một hàng / câu)."""
    keys = ("a", "b", "c", "d")
    vals = [it.get(k) if it.get(k) is not None else it.get(k.upper(), "") for k in keys]
    ans = str(it.get("ans", "A")).strip().upper()[:1]
    ai = ord(ans) - ord("A")
    if ai not in range(4):
        return it
    order = [0, 1, 2, 3]
    random.shuffle(order)
    new_vals = [vals[order[j]] for j in range(4)]
    new_ai = order.index(ai)
    new_ans = chr(ord("A") + new_ai)
    out = {**it, "a": new_vals[0], "b": new_vals[1], "c": new_vals[2], "d": new_vals[3], "ans": new_ans}
    return out


def _parse_vertical_quiz_items(arr: list[Any]) -> list[VerticalQuizBlock]:
    if len(arr) != 5:
        raise ValueError(f"Cần đúng 5 câu (khung cố định), model trả {len(arr)} phần tử.")
    blocks: list[VerticalQuizBlock] = []
    for idx, item in enumerate(arr):
        if not isinstance(item, dict):
            raise ValueError(f"Câu {idx + 1}: mỗi phần tử phải là object JSON.")
        stt, md, mt = QUIZ_FIXED_LEVELS[idx]
        qv = _flex_question_text(item)
        if not qv:
            raise ValueError(
                f"Câu {idx + 1}: thiếu nội dung câu hỏi (khóa hợp lệ: q, question, …)."
            )
        opts = _flex_options_list(item)
        if not isinstance(opts, list) or len(opts) != 4:
            raise ValueError(f"Câu {idx + 1}: cần đúng 4 đáp án trong mảng options/choices.")
        row_choices: list[tuple[str, str, str]] = []
        n_ok = 0
        for j, o in enumerate(opts):
            if not isinstance(o, dict):
                raise ValueError(f"Câu {idx + 1}, đáp án {j + 1}: không phải object.")
            t = _flex_opt_text(o)
            e = _flex_opt_explain(o)
            okb = _flex_opt_ok(o)
            if t is None:
                raise ValueError(f"Câu {idx + 1}, đáp án {j + 1}: thiếu nội dung (t/text).")
            if okb is True:
                n_ok += 1
            elif okb is False:
                pass
            else:
                raise ValueError(
                    f"Câu {idx + 1}, đáp án {j + 1}: thiếu ok/correct (boolean)."
                )
            e_clean = clean_quiz_explanation(e)
            if not e_clean or len(e_clean) < 3:
                raise ValueError(
                    f"Câu {idx + 1}, đáp án {j + 1}: mỗi đáp án phải có giải thích riêng "
                    f"(e, sau khi bỏ tiền tố Đúng/Sai còn ít nhất 3 ký tự)."
                )
            kq = "Đúng" if okb is True else "Sai"
            row_choices.append((t, kq, e_clean))
        if n_ok != 1:
            raise ValueError(f"Câu {idx + 1}: phải đúng một đáp án đúng (ok:true), đang có {n_ok}.")
        blocks.append(
            VerticalQuizBlock(
                stt=stt,
                muc_do=md,
                muc_tieu=mt,
                question_text=qv,
                choices=(row_choices[0], row_choices[1], row_choices[2], row_choices[3]),
            )
        )
    return blocks


def run_quiz_generation(params: QuizGenParams) -> tuple[bool, str]:
    """
    Trả về (thành công, thông báo lỗi hoặc đường dẫn file đã ghi).
    """
    if not params.template_xlsx.is_file():
        return False, f"Không tìm thấy file Excel mẫu: {params.template_xlsx}"

    qkind = normalize_quiz_kind(params.quiz_kind)
    if qkind in (QUIZ_KIND_SESSION_WARMUP, QUIZ_KIND_SESSION_END):
        if not (params.subject or "").strip():
            return False, "Thiếu môn học."
        if not (params.session_current or "").strip():
            return False, "Thiếu session hiện tại."
    else:
        if not (params.lesson or "").strip():
            return False, "Thiếu tên lesson."
        if not (params.session or "").strip():
            return False, "Thiếu tên session."

    # Tài liệu bài giảng: ưu tiên text (Google Docs), fallback DOCX (cũ).
    lecture_text = (params.lecture_text or "").strip()
    docx_text = ""
    if not lecture_text and params.docx_path:
        p = Path(params.docx_path)
        if not p.is_file():
            return False, f"Không tìm thấy file DOCX: {p}"
        try:
            docx_text = _extract_docx_plain(p)
        except Exception as e:
            return False, f"Lỗi đọc DOCX: {e}"
        lecture_text = (docx_text or "").strip()

    # Quizz session đầu giờ / cuối giờ: 45 câu, 1 hàng / câu theo header warmup
    if qkind in (QUIZ_KIND_SESSION_WARMUP, QUIZ_KIND_SESSION_END):
        try:
            headers = read_headers_from_template(params.template_xlsx)
        except Exception as e:
            return False, f"Không đọc được Excel mẫu: {e}"
        if not headers:
            return False, "File mẫu không có tiêu đề cột ở hàng 1."

        subj = (params.subject or "").strip()
        prev_s = (params.session_prev or "").strip()
        curr_s = (params.session_current or "").strip()
        docx_excerpt = lecture_text[:9000] if lecture_text.strip() else ""
        m = resolve_model(params.model or None)
        temp0 = max(0.0, min(2.0, float(params.temperature)))

        # Chia 45 câu thành 3 block 15 câu để giảm lỗi cắt/thiếu JSON và thường nhanh hơn.
        blocks: list[tuple[str, int, int]]
        if qkind == QUIZ_KIND_SESSION_WARMUP:
            blocks = [("prev", 1, 15), ("prev", 16, 15), ("current", 31, 15)]
        else:
            blocks = [("current", 1, 15), ("current", 16, 15), ("current", 31, 15)]
        all_items: list[Any] = []
        for part, start_stt, n_need in blocks:
            last_raw = ""
            arr_block: list[Any] | None = None
            for attempt in range(3):
                if qkind == QUIZ_KIND_SESSION_WARMUP:
                    if attempt == 0:
                        msgs = _warmup_block_messages(
                            subject=subj,
                            session_prev=prev_s,
                            session_current=curr_s,
                            part=part,
                            start_stt=start_stt,
                            n=n_need,
                            docx_excerpt=docx_excerpt,
                        )
                    else:
                        msgs = _warmup_block_retry_messages(bad_raw=last_raw, n=n_need, part=part)
                else:
                    if attempt == 0:
                        msgs = _end_block_messages(
                            subject=subj,
                            session_current=curr_s,
                            start_stt=start_stt,
                            n=n_need,
                            docx_excerpt=docx_excerpt,
                        )
                    else:
                        msgs = _end_block_retry_messages(bad_raw=last_raw, n=n_need)
                raw, _data = complete_chat(
                    msgs,
                    model=m,
                    temperature=min(temp0, 0.35),
                    max_tokens=9000,
                    timeout_s=420.0,
                )
                last_raw = raw
                try:
                    arr_block = _parse_json_array(raw)
                    if len(arr_block) != n_need:
                        raise ValueError(f"Cần {n_need} câu, nhận {len(arr_block)}.")
                    break
                except Exception:
                    if attempt >= 2:
                        arr_block = None
                        break
            if arr_block is None:
                if qkind == QUIZ_KIND_SESSION_WARMUP:
                    return (
                        False,
                        f"JSON warmup bị cắt/sai ở block {start_stt}–{start_stt + n_need - 1}.\n---\n{last_raw[:2800]}",
                    )
                return (
                    False,
                    f"JSON session cuối giờ bị cắt/sai ở block {start_stt}–{start_stt + n_need - 1}.\n---\n{last_raw[:2800]}",
                )
            all_items.extend(arr_block)

        try:
            rows = _parse_session_warmup_items(all_items, plan=("warmup" if qkind == QUIZ_KIND_SESSION_WARMUP else "end"))
        except Exception as e:
            if qkind == QUIZ_KIND_SESSION_WARMUP:
                return False, f"JSON warmup không hợp lệ: {e}\n---\n{json.dumps(all_items[:2], ensure_ascii=False)[:2800]}"
            return (
                False,
                f"JSON session cuối giờ không hợp lệ: {e}\n---\n{json.dumps(all_items[:2], ensure_ascii=False)[:2800]}",
            )

        fill_template_session_warmup_quiz(params.template_xlsx, params.output_xlsx, rows)
        return True, str(params.output_xlsx)

    try:
        headers = read_headers_from_template(params.template_xlsx)
    except Exception as e:
        return False, f"Không đọc được Excel mẫu: {e}"

    if not headers:
        return False, "File mẫu không có tiêu đề cột ở hàng 1."

    vertical = is_vertical_quiz_template(headers)
    n = 5 if vertical else max(1, min(50, int(params.num_questions)))

    if vertical:
        spec_lines = "\n".join(
            f"  Câu {stt}: Mức độ = «{md}», Mục tiêu = «{mt}» (bạn PHẢI soạn nội dung câu hỏi và đáp án sát "
            f"đúng cặp mức độ/mục tiêu này)"
            for stt, md, mt in QUIZ_FIXED_LEVELS
        )
        lesson_line = f"- Lesson: {params.lesson.strip()}\n"
        session_line = f"- Session: {params.session.strip()}\n"
        if qkind == QUIZ_KIND_SESSION_WARMUP:
            lesson_line = ""
            session_line = ""
        user = f"""Thông tin:
{lesson_line}{session_line}"""
        if qkind == QUIZ_KIND_SESSION_WARMUP:
            user += (
                f"- Session trước: {(params.session_prev or '').strip()}\n"
                f"- Session hiện tại: {(params.session_current or '').strip()}\n"
            )
        user += f"""{_quiz_kind_bullet(qkind)}
{_subject_bullet(params.subject)}
Khung 5 câu — trên Excel, cột «Mức độ» đúng 5 nhãn (đã cố định, không đổi tên):
  Câu 1–2: Thông hiểu | Câu 3: Vận dụng sơ bộ | Câu 4: Phân biệt/So sánh | Câu 5: Phân tích sơ bộ
  (Không dùng «Nhớ», «Hiểu», «Vận dụng» không có «sơ bộ» — mức độ do tool ghi, bạn chỉ bám mục tiêu từng dòng.)

Chi tiết mục tiêu & mức độ từng câu:
{spec_lines}
"""
        if lecture_text.strip():
            user += f"""
Nội dung bài giảng (tài liệu cung cấp):
---
{lecture_text[:12000]}
---
"""
        else:
            if qkind == QUIZ_KIND_SESSION_WARMUP:
                user += (
                    "\n(Không có DOCX — soạn quiz sát session hiện tại; cho phép 1–2 câu nhắc lại kiến thức session trước.)\n"
                )
            else:
                user += (
                    "\n(Không có DOCX — hãy soạn quiz sát lesson/session và đúng mức độ/mục tiêu từng câu.)\n"
                )
        user += (
            "\nTrả về đúng mảng JSON 5 phần tử theo schema system (khóa ASCII q/options/t/ok/e). "
            "Mỗi \"e\" chỉ lý do (không ghi \"Đúng:\" / \"Sai:\"). "
            "Nếu có code Python trong \"q\": tên biến/hàm chỉ tiếng Anh; phải đúng logic và khớp đáp án — không dùng continue/if kiểu i chẵn lẻ làm thiếu giá trị (vd. chỉ in 1,3,5 mà hỏi có 4). "
            "Kết thúc bằng ].\n"
        )
        system_msg = SYSTEM_QUIZ_VERTICAL
    else:
        user = "Thông tin:\n"
        if qkind != QUIZ_KIND_SESSION_WARMUP:
            user += f"- Lesson: {params.lesson.strip()}\n- Session: {params.session.strip()}\n"
        else:
            user += (
                f"- Session trước: {(params.session_prev or '').strip()}\n"
                f"- Session hiện tại: {(params.session_current or '').strip()}\n"
            )
        user += f"""{_subject_bullet(params.subject)}{_quiz_kind_bullet(qkind)}- Số câu hỏi: {n}

Định dạng JSON xem system: mỗi câu là object với khóa stt, q, a, b, c, d, ans (chỉ ASCII).
Chương trình sẽ tự ghép vào đúng cột Excel mẫu — không dùng tên cột tiếng Việt làm khóa JSON.
"""
        if lecture_text.strip():
            user += f"""
Nội dung bài giảng (tài liệu cung cấp):
---
{lecture_text[:12000]}
---
"""
        else:
            if qkind == QUIZ_KIND_SESSION_WARMUP:
                user += "\n(Không có DOCX — soạn quiz sát session hiện tại; cho phép 1–2 câu nhắc lại session trước.)\n"
            else:
                user += "\n(Không có DOCX — soạn quiz sát lesson/session.)\n"

        user += (
            f"\nTrả về đúng {n} phần tử trong mảng JSON. "
            "Nếu có code trong \"q\": tên biến/hàm chỉ tiếng Anh.\n"
        )
        system_msg = SYSTEM_QUIZ

    m = resolve_model(params.model or None)
    mt_user = max(256, int(params.max_tokens))
    max_tok = max(mt_user, 32768 if vertical else 16384)
    temp0 = max(0.0, min(2.0, float(params.temperature)))
    raw, data = complete_chat(
        [
            ChatMessage(role="system", content=system_msg),
            ChatMessage(role="user", content=user),
        ],
        model=m,
        temperature=temp0,
        max_tokens=max_tok,
    )

    arr: list[Any] | None = None
    parse_err: Exception | None = None
    try:
        arr = _parse_json_array(raw)
    except Exception as e:
        parse_err = e

    fr = _finish_reason(data)
    if (arr is None or fr == "length") and vertical:
        try:
            raw2, _ = complete_chat(
                _vertical_retry_messages(raw),
                model=m,
                temperature=0.05,
                max_tokens=max_tok,
            )
            arr = _parse_json_array(raw2)
            parse_err = None
            raw = raw2
        except Exception as e2:
            if parse_err is None:
                parse_err = e2
    elif (arr is None or fr == "length") and not vertical:
        try:
            raw2, _ = complete_chat(
                _flat_retry_messages(raw, n),
                model=m,
                temperature=0.05,
                max_tokens=max_tok,
            )
            arr = _parse_json_array(raw2)
            parse_err = None
            raw = raw2
        except Exception as e2:
            if parse_err is None:
                parse_err = e2

    if arr is None:
        hint = ""
        if fr == "length":
            hint = (
                "\n(Gợi ý: output bị cắt do max_tokens — đã tăng hạn mức và thử gọi lại; "
                "nếu vẫn lỗi, đổi model hoặc rút ngắn DOCX.)"
            )
        return False, f"Model không trả JSON hợp lệ: {parse_err}{hint}\n---\n{raw[:2800]}"

    if not isinstance(arr, list):
        return False, "JSON phải là mảng các câu hỏi."

    try:
        if vertical:
            blocks = _parse_vertical_quiz_items(arr)
            slots = _random_correct_slots_five()
            blocks = [
                VerticalQuizBlock(
                    b.stt,
                    b.muc_do,
                    b.muc_tieu,
                    b.question_text,
                    _reorder_vertical_choices_to_slot(b.choices, slots[i]),
                )
                for i, b in enumerate(blocks)
            ]
            fill_template_vertical_quiz(params.template_xlsx, params.output_xlsx, blocks)
        else:
            if len(arr) < n:
                return False, f"Model chỉ trả {len(arr)} câu, cần {n}."
            arr = arr[:n]
            shuffled = []
            for it in arr:
                if isinstance(it, dict):
                    shuffled.append(_shuffle_flat_item_answers(it))
                else:
                    shuffled.append(it)
            rows = flat_english_items_to_rows(shuffled, headers)
            fill_template_from_rows(params.template_xlsx, params.output_xlsx, rows)
    except Exception as e:
        return False, f"Lỗi ghi Excel: {e}"

    return True, str(Path(params.output_xlsx).resolve())
