from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from cham_bai.collector import DEFAULT_IGNORE_DIR_NAMES, CollectedBundle

GITHUB_URL_RE = re.compile(
    r"https://github\.com/(?P<owner>[\w.-]+)/(?P<repo>[\w.-]+)(?:/(?:tree|blob)/[\w./-]+)?/?",
    re.IGNORECASE,
)
GIT_SSH_RE = re.compile(
    r"git@github\.com:(?P<owner>[\w.-]+)/(?P<repo>[\w.-]+?)(?:\.git)?$",
    re.IGNORECASE,
)


@dataclass
class DocxContent:
    plain_text: str
    github_repo_urls: list[str]


def _normalize_repo_url(owner: str, repo: str) -> str:
    owner, repo = owner.strip(), repo.strip().rstrip("/")
    if repo.endswith(".git"):
        repo = repo[:-4]
    return f"https://github.com/{owner}/{repo}"


def _urls_from_word_hyperlinks(doc: Document) -> list[str]:
    out: list[str] = []
    try:
        for rel in doc.part.rels.values():
            if rel.reltype != RT.HYPERLINK:
                continue
            t = (rel.target_ref or "").strip()
            if t:
                out.append(t)
    except (KeyError, AttributeError, TypeError):
        pass
    return out


def github_urls_from_plain_text(*blobs: str) -> list[str]:
    """Thu thập link GitHub từ một hoặc nhiều chuỗi (đoạn văn, URL...)."""
    seen: set[str] = set()
    urls: list[str] = []
    for blob in blobs:
        _register_github_urls(blob, seen, urls)
    return urls


def _register_github_urls(blob: str, seen: set[str], urls: list[str]) -> None:
    for m in GITHUB_URL_RE.finditer(blob):
        url = _normalize_repo_url(m.group("owner"), m.group("repo"))
        if url not in seen:
            seen.add(url)
            urls.append(url)
    for m in GIT_SSH_RE.finditer(blob.strip()):
        url = _normalize_repo_url(m.group("owner"), m.group("repo"))
        if url not in seen:
            seen.add(url)
            urls.append(url)


def _document_to_docx_content(doc: Document) -> DocxContent:
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    plain = "\n".join(parts)
    seen: set[str] = set()
    urls: list[str] = []
    for blob in [plain, *_urls_from_word_hyperlinks(doc)]:
        _register_github_urls(blob, seen, urls)

    return DocxContent(plain_text=plain, github_repo_urls=urls)


def extract_docx(path: str) -> DocxContent:
    return _document_to_docx_content(Document(path))


def extract_docx_bytes(data: bytes) -> DocxContent:
    """Đọc .docx từ bytes (ví dụ tải từ OneDrive)."""
    return _document_to_docx_content(Document(io.BytesIO(data)))


def _docx_rel_should_skip(rel: Path) -> bool:
    for part in rel.parts:
        if part in DEFAULT_IGNORE_DIR_NAMES:
            return True
    return False


def append_docx_plaintext_from_repo_to_bundle(
    root: Path,
    bundle: CollectedBundle,
    *,
    max_docx: int = 8,
    max_chars_each: int = 90_000,
    out_warnings: list[str] | None = None,
) -> None:
    """
    Đọc các file .docx dưới root (repo đã clone), thêm nội dung văn bản vào bundle.files
    (collector bỏ qua .docx vì nhị phân). Gọi trước khi xóa thư mục tạm clone.
    """
    n = 0
    for p in sorted(root.rglob("*.docx")):
        if not p.is_file():
            continue
        try:
            rel = p.relative_to(root)
        except ValueError:
            continue
        if _docx_rel_should_skip(rel):
            continue
        if n >= max_docx:
            if out_warnings is not None:
                out_warnings.append(
                    f"Chỉ trích tối đa {max_docx} file DOCX trong repo; bỏ qua các file .docx còn lại."
                )
            break
        try:
            dc = extract_docx(str(p))
        except Exception as e:
            if out_warnings is not None:
                out_warnings.append(f"Không đọc được {rel.as_posix()}: {e}")
            continue
        text = (dc.plain_text or "").strip()
        if not text:
            continue
        if len(text) > max_chars_each:
            text = text[:max_chars_each] + "\n\n[... DOCX truncated ...]\n"
        pseudo = f"_docx_text/{rel.as_posix()}.txt"
        bundle.files.append((pseudo, text))
        n += 1
