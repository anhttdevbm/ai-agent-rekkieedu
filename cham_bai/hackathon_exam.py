from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@dataclass
class HackathonExamParams:
    header_top: str  # e.g. "KIỂM TRA HACKATHON"
    header_sub: str  # e.g. "NHẬP MÔN CSDL MYSQL - Đề 006"
    duration_minutes: int
    body_text: str


_BONUS_NOTE_1 = "Bonus điểm : clean code, đầy đủ comment, trình bày đẹp cộng tối đa 5 điểm."
_BONUS_NOTE_2 = "Lưu ý: Chỉ tính điểm khi thực hiện đúng theo yêu cầu"


def build_hackathon_exam_docx_from_spec(spec: dict[str, Any]) -> bytes:
    """
    Render DOCX from a structured JSON spec (AI output).
    Expected keys:
    - header_top, header_sub, duration_minutes
    - sections: [
        { "title": str,
          "paragraphs": [str],
          "bullets": [str],
          "numbered": [str],
          "tables": [
             { "title": str?, "headers": [str], "rows": [[str,...], ...] }
          ]
        }, ...
      ]
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    header_top = str(spec.get("header_top") or "KIỂM TRA HACKATHON").strip()
    header_sub = str(spec.get("header_sub") or "").strip()
    mins = int(spec.get("duration_minutes") or 120)

    _add_center(doc, header_top, bold=True, size=16)
    if header_sub:
        _add_center(doc, header_sub, bold=True, size=14)
    _add_center(doc, f"THỜI GIAN: {max(1, mins)} phút", bold=True, size=13)
    _add_center(doc, _STAR_LINE, bold=False, size=12)
    doc.add_paragraph("")

    sections = spec.get("sections") or []
    if not isinstance(sections, list):
        sections = []

    for sec in sections:
        if not isinstance(sec, dict):
            continue
        title = str(sec.get("title") or "").strip()
        if title:
            p = doc.add_paragraph(title)
            if p.runs:
                p.runs[0].bold = True

        paras = sec.get("paragraphs") or []
        if isinstance(paras, list):
            for t in paras:
                s = str(t or "").strip()
                if s:
                    doc.add_paragraph(s)

        bullets = sec.get("bullets") or []
        if isinstance(bullets, list):
            for t in bullets:
                s = str(t or "").strip()
                if s:
                    doc.add_paragraph(s, style="List Bullet")

        tables = sec.get("tables") or []
        if isinstance(tables, list):
            for tb in tables:
                if not isinstance(tb, dict):
                    continue
                tb_title = str(tb.get("title") or "").strip()
                if tb_title:
                    p2 = doc.add_paragraph(tb_title)
                    if p2.runs:
                        p2.runs[0].bold = True
                headers = tb.get("headers") or []
                rows = tb.get("rows") or []
                if not isinstance(headers, list) or not headers:
                    continue
                if not isinstance(rows, list):
                    rows = []
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h or "")
                for r in rows[:500]:
                    if not isinstance(r, list):
                        continue
                    cells = table.add_row().cells
                    for i in range(len(headers)):
                        cells[i].text = str(r[i] if i < len(r) else "")
                doc.add_paragraph("")

        # Plain lines (no bullets/no Word numbering).
        # Important: lines must appear AFTER tables so questions sit right under tables.
        lines = sec.get("lines") or []
        if isinstance(lines, list):
            for t in lines:
                s = str(t or "").strip()
                if s:
                    doc.add_paragraph(s)

        # Back-compat: if older specs still produce "numbered", render as plain lines
        # (NOT Word "List Number", to avoid auto-renumber/reset).
        numbered = sec.get("numbered") or []
        if isinstance(numbered, list):
            for t in numbered:
                s = str(t or "").strip()
                if s:
                    doc.add_paragraph(s)

    doc.add_paragraph("")
    doc.add_paragraph(_BONUS_NOTE_1)
    doc.add_paragraph(_BONUS_NOTE_2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_STAR_LINE = "*" * 19


def _clean_lines(s: str) -> list[str]:
    raw = (s or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.rstrip() for ln in raw.split("\n")]
    # trim empty head/tail
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def _add_center(doc: Document, text: str, *, bold: bool = False, size: int = 14) -> None:
    p = doc.add_paragraph(text or "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0] if p.runs else p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)


def _add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph(text or "")
    if bold and p.runs:
        p.runs[0].bold = True


_PART_RE = re.compile(r"^\s*PHẦN\s+\d+\s*:", re.I)
_SECTION_RE = re.compile(r"^\s*\d+\.\s+")
_RUBRIC_HEADER_RE = re.compile(r"^\s*Thang\s+chấm\s+điểm\s*:?\s*$", re.I)
_RUBRIC_ROW_RE = re.compile(r"^\s*(?P<q>\d+)\s+(?P<content>.+?)\s+(?P<pts>\d+)\s*$")


def _iter_rubric_rows(lines: Iterable[str]) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    for ln in lines:
        m = _RUBRIC_ROW_RE.match(ln)
        if not m:
            continue
        q = m.group("q").strip()
        content = m.group("content").strip()
        pts = m.group("pts").strip()
        rows.append((q, content, pts))
    return rows


def build_hackathon_exam_docx_bytes(params: HackathonExamParams) -> bytes:
    doc = Document()

    # base font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    header_top = (params.header_top or "").strip()
    header_sub = (params.header_sub or "").strip()
    mins = max(1, int(params.duration_minutes or 0))

    _add_center(doc, header_top or "KIỂM TRA HACKATHON", bold=True, size=16)
    if header_sub:
        _add_center(doc, header_sub, bold=True, size=14)
    _add_center(doc, f"THỜI GIAN: {mins} phút", bold=True, size=13)
    _add_center(doc, _STAR_LINE, bold=False, size=12)
    doc.add_paragraph("")  # spacing

    lines = _clean_lines(params.body_text)

    i = 0
    while i < len(lines):
        ln = lines[i].strip()
        if not ln:
            doc.add_paragraph("")
            i += 1
            continue

        if _PART_RE.match(ln) or re.match(r"^\s*\d+\.\s*[^:]+:\s*$", ln):
            p = doc.add_paragraph(ln)
            p.runs[0].bold = True
            i += 1
            continue

        # bullets (• or -)
        if ln.startswith("•") or ln.startswith("- "):
            text = ln.lstrip("•").lstrip("-").strip()
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # numbered item "1." "2." etc
        if _SECTION_RE.match(ln):
            # render as plain line to preserve 1..N numbering written in text
            doc.add_paragraph(ln)
            i += 1
            continue

        _add_paragraph(doc, ln)
        i += 1

    # Bonus note at the end (like mẫu)
    doc.add_paragraph("")
    doc.add_paragraph(_BONUS_NOTE_1)
    doc.add_paragraph(_BONUS_NOTE_2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

