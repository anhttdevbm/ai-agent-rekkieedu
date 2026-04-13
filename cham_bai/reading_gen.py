from __future__ import annotations

import base64
import io
import keyword
import re
import tokenize
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from cham_bai.openrouter import ChatMessage, complete_chat, generate_images_from_prompt
from cham_bai.settings import model as resolve_model
from cham_bai.video_transcript import fetch_youtube_transcript_plain

# Thứ tự ưu tiên khi model ảnh người chọn không trả về `message.images`.
READING_IMAGE_MODEL_FALLBACKS: tuple[str, ...] = (
    "black-forest-labs/flux.2-pro",
    "google/gemini-2.5-flash-image-preview",
    "google/gemini-2.5-flash-image",
    "black-forest-labs/flux.2-flex",
)


def _reading_first_image_blob(prompt: str, image_model_primary: str) -> bytes | None:
    """Một ảnh đầu tiên từ prompt; thử lần lượt primary + fallback models."""
    for m_try in _reading_image_models_to_try(image_model_primary):
        blobs = generate_images_from_prompt(prompt, model=m_try)
        if blobs:
            return blobs[0]
    return None


def _reading_image_models_to_try(primary: str) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for m in ((primary or "").strip(),) + READING_IMAGE_MODEL_FALLBACKS:
        if not m or m in seen:
            continue
        seen.add(m)
        out.append(m)
    return out


DEFAULT_LEARNING_GOALS = """Sau bài này, sinh viên có thể:
• Hiểu bản chất của khái niệm/cú pháp trọng tâm
• Biết khi nào nên dùng trong tình huống thực tế
• Đọc và phân tích được code mẫu
• Áp dụng để giải quyết một bài toán cùng domain"""

# Nhân vật ảo minh họa mục I (tiếng Anh cho model ảnh): 3D CGI điện ảnh siêu thực; không phải ảnh chụp người thật / celebrity.
READING_VIRTUAL_PROTAGONIST_MALE_EN = (
    "Hyper-realistic cinematic 3D CGI protagonist (NOT a real photograph, NOT identifiable real person): recurring young adult "
    "man, male protagonist only, East Asian ~early twenties, oval face with natural micro-detail and subsurface glow on skin, "
    "black hair in soft layered cut with mild mushroom / rounded-bowl volume over the crown (not severe bowl cut), "
    "short choppy mushroom bangs covering the forehead (playful blunt micro-bangs vibe), gentle loose waves or light perm "
    "curl, lightly tousled bedhead texture; NO visible center or side part; full sides and back—NO undercut, NO faded or "
    "shaved sides; monolid dark brown eyes, straight dark eyebrows, soft jawline; expression "
    "can match story—calm professional focus, subtle stress or concern at a workstation crisis when appropriate; wearing the "
    "SAME outfit always: crisp white dress shirt with collar, hem worn UNTUCKED (no tuck-in), slightly relaxed "
    "or oversized Gen Z smart-street vibe; shirt BUTTONED CLOSED down the chest—no open placket, no deep unbuttoned chest; "
    "at most only the TOP one OR two collar buttons left open; light-blue or whitewash denim jeans with ripped knees (distressed knee holes), "
    "casual urban look; moderate build; cinematic cool-toned lighting, soft shadows "
    "sculpting the face, shallow depth of field, dark blurred bokeh background, high-end game-cinematic or feature-quality "
    "3D render (Unreal Engine 5 style acceptable), 8k detail mood, consistent same avatar across images, generic "
    "non-celebrity digital character."
)

READING_VIRTUAL_PROTAGONIST_FEMALE_EN = (
    "Hyper-realistic cinematic 3D CGI heroine (NOT a real photograph, NOT identifiable real person): recurring young adult "
    "woman, female protagonist only, East Asian late teens / early twenties, soft rounded face, fair skin with natural "
    "micro-detail and subsurface glow, long straight dark brown or black hair with clean center part, dark eyes, defined "
    "dark brows, deep pink or berry lipstick; expression can match story—subtle worry, concern, or calm focus at a "
    "workstation crisis when appropriate (readable micro-expression, slightly glassy eyes if tense); wearing the SAME "
    "outfit always: white strapless off-the-shoulder tube dress with dramatic tiered ruffled layers of soft sheer or "
    "organza-like fabric across bust and torso, shoulders bare, very thin delicate gold necklace, thin bracelet on wrist; "
    "moderate build; cinematic cool-toned lighting, soft shadows sculpting the face, shallow depth of field, dark blurred "
    "bokeh background, high-end game-cinematic or feature-quality 3D render (Unreal Engine 5 style acceptable), 8k detail "
    "mood, consistent same avatar across images, generic non-celebrity digital character."
)


def sanitize_reading_filename_part(s: str, max_len: int = 56) -> str:
    s = (s or "").strip()
    if not s:
        return "reading_doc"
    s = re.sub(r"[\x00-\x1f]", "", s)
    s = re.sub(r'[<>:"/\\|?*]', "_", s)
    s = re.sub(r"\s+", "_", s)
    s = s.strip("._") or "reading_doc"
    return s[:max_len]


def _sanitize_subject_for_title_filename(s: str) -> str:
    """Tên môn hiển thị trên file: bỏ ký tự cấm Windows, gộp khoảng trắng."""
    s = (s or "").strip()
    s = re.sub(r"[\x00-\x1f]", "", s)
    s = re.sub(r'[<>:"/\\|?*]', "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s or "Subject"


def reading_output_stem(subject: str, session_stt: str, lesson_stt: str, *, max_total: int = 200) -> str:
    """
    Quy tắc: «Tên môn học - Session x - Lesson y» (x, y là số thứ tự người nhập).
    Rút gọn tên môn nếu quá dài để tránh đường dẫn Windows.
    """
    ss = str(session_stt).strip() or "1"
    ls = str(lesson_stt).strip() or "1"
    suffix = f" - Session {ss} - Lesson {ls}"
    sub = _sanitize_subject_for_title_filename(subject)
    budget = max_total - len(suffix)
    if len(sub) > budget:
        sub = sub[: max(budget - 1, 12)].rstrip() + "…"
    return f"{sub}{suffix}"


def default_reading_output_pair(stem: str) -> tuple[Path, Path]:
    """Hai đường dẫn .docx và .xlsx cùng stem, tại thư mục hiện tại; tránh trùng tên."""
    base = Path.cwd()
    for n in range(0, 500):
        extra = f" ({n})" if n else ""
        docx_p = base / f"{stem}{extra}.docx"
        xlsx_p = base / f"{stem}{extra}.xlsx"
        if not docx_p.exists() and not xlsx_p.exists():
            return docx_p, xlsx_p
    raise RuntimeError("Không tìm được tên file trống.")


def default_reading_output_path(title_hint: str) -> Path:
    """Tương thích cũ: chỉ trả về docx (dùng khi chưa có stem chuẩn)."""
    stem = f"reading_doc_{sanitize_reading_filename_part(title_hint)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    docx_p, _ = default_reading_output_pair(stem)
    return docx_p


def split_markdown_remove_vii_for_docx(md: str) -> tuple[str, str | None]:
    """
    Tách mục ## VII… để DOCX không in; phần VII giữ làm nguồn Excel.
    Trả về (markdown cho DOCX: I–VI rồi **## VIII.** tài liệu — trước khi in sẽ đổi VIII→VII),
    và nội dung mục câu hỏi (gồm dòng ## VII.) hoặc None.
    """
    lines = md.splitlines(keepends=True)
    vii_start: int | None = None
    viii_start: int | None = None
    for i, line in enumerate(lines):
        s = line.strip()
        if re.match(r"^##\s+VII[\.\s]", s):
            vii_start = i
        elif vii_start is not None and re.match(r"^##\s+VIII[\.\s]", s):
            viii_start = i
            break
    if vii_start is None:
        return md, None
    doc_lines = lines[:vii_start]
    if viii_start is not None:
        doc_lines.extend(lines[viii_start:])
    vii_lines = lines[vii_start:viii_start if viii_start is not None else len(lines)]
    return "".join(doc_lines), "".join(vii_lines).strip()


_HR_ONLY_LINE = re.compile(r"^\s*(?:-{3,}|_{3,}|\*{3,})\s*$")


def _strip_trailing_horizontal_rules(text: str) -> str:
    """Bỏ các dòng chỉ có --- / *** / ___ (Markdown HR) ở cuối — model hay thêm sau Đáp án."""
    t = text.rstrip()
    while True:
        lines = t.splitlines()
        if not lines:
            return t
        if _HR_ONLY_LINE.match(lines[-1]):
            t = "\n".join(lines[:-1]).rstrip()
            continue
        return t


def strip_markdown_for_excel_cell(text: str) -> str:
    """
    Gỡ dấu hiệu Markdown khỏi ô Excel: fence ```, và bọc in đậm **…**.
    Giữ nguyên phép lũy thừa Python dạng «số ** số» (ví dụ 5 ** 2).
    """
    if not text:
        return text
    shielded: list[str] = []

    def _shield_exp(m: re.Match[str]) -> str:
        shielded.append(m.group(0))
        return f"\x00EXP{len(shielded) - 1}\x00"

    t = re.sub(r"\d\s*\*\*\s*\d", _shield_exp, text)
    t = re.sub(r"(?m)^\s*```[\w]*\s*$", "", t)
    t = t.replace("```", "")
    # Biến/placeholder: {{n}} / {{n-1}} -> n / n-1 (để Excel sạch, không lộ marker)
    t = _MD_INLINE_VAR.sub(lambda m: re.sub(r"\s+", "", m.group(1)), t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    prev = None
    while prev != t:
        prev = t
        t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = t.replace("**", "")
    for i, chunk in enumerate(shielded):
        t = t.replace(f"\x00EXP{i}\x00", chunk)
    t = t.strip()
    t = _strip_trailing_horizontal_rules(t)
    return t


def renumber_references_viii_to_vii_for_docx(md: str) -> str:
    """
    Sau khi bỏ mục câu hỏi (VII trong bản thảo), phần tài liệu vẫn là ## VIII. trong Markdown
    nhưng trên DOCX phải là mục **VII** cho đúng thứ tự I…VI → VII.
    Chỉ đổi **lần xuất hiện đầu tiên** của tiêu đề ## VIII (tránh đụng nhầm nội dung code).
    """
    return re.sub(r"(^##\s+)VIII([\.\s])", r"\1VII\2", md, count=1, flags=re.MULTILINE)


_VII_CQ = re.compile(
    r"(?im)(?:^|\r?\n)\s*\*{0,2}\s*Câu\s*hỏi\s*:\s*\*{0,2}\s*",
)
_VII_DA = re.compile(
    r"(?im)(?:^|\r?\n)\s*\*{0,2}\s*Đáp\s*án\s*:\s*\*{0,2}\s*",
)


def _split_vii_block_question_answer(block: str) -> tuple[str, str]:
    """Tách **Câu hỏi:** … **Đáp án:** … trong một khối (vd. sau nhãn Thông hiểu)."""
    block = (block or "").strip()
    if not block:
        return "", ""
    mq = _VII_CQ.search(block)
    ma = _VII_DA.search(block)
    if mq and ma and ma.start() > mq.end():
        q = block[mq.end() : ma.start()].strip()
        a = block[ma.end() :].strip()
        return q, a
    if mq and not ma:
        return block[mq.end() :].strip(), ""
    if not mq and ma:
        return block[: ma.start()].strip(), block[ma.end() :].strip()
    return block, ""


def _parse_vii_qa_by_labels_only(t: str) -> list[tuple[str, str]]:
    """Lần lượt các cặp Câu hỏi / Đáp án trong toàn bộ mục VII (khi cần fallback)."""
    out: list[tuple[str, str]] = []
    pos = 0
    while len(out) < 3:
        mq = _VII_CQ.search(t, pos)
        if not mq:
            break
        ma = _VII_DA.search(t, mq.end())
        if not ma:
            out.append((t[mq.end() :].strip(), ""))
            break
        mq_next = _VII_CQ.search(t, ma.end())
        end_a = mq_next.start() if mq_next else len(t)
        q = t[mq.end() : ma.start()].strip()
        a = t[ma.end() : end_a].strip()
        out.append((q, a))
        pos = end_a
    return out


def _parse_vii_excel_qa_rows(section_vii_markdown: str) -> list[tuple[str, str]]:
    """Tối đa 3 cặp (questionName, answerName) từ Markdown mục VII."""
    t = (section_vii_markdown or "").strip()
    t = re.sub(r"^##\s+VII\.\s*[^\n]*\n?", "", t, count=1, flags=re.M).strip()
    if not t:
        return []
    pattern = re.compile(
        r"(?is)(?:^|\n)\s*(?:#{1,3}\s*|\*\*)*\s*"
        r"(Thông\s*hiểu|Vận\s*dụng|Phân\s*tích)"
        r"\s*(?:\*\*)?\s*[:\.]?\s*",
    )
    matches = list(pattern.finditer(t))
    pairs: list[tuple[str, str]] = []
    if matches:
        for i, m in enumerate(matches):
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(t)
            body = t[start:end].strip()
            if body:
                pairs.append(_split_vii_block_question_answer(body))
        pairs = pairs[:3]
    if not pairs:
        pairs = _parse_vii_qa_by_labels_only(t)
    elif pairs and not any(a.strip() for _, a in pairs) and _VII_DA.search(t):
        alt = _parse_vii_qa_by_labels_only(t)
        if alt and any(a.strip() for _, a in alt):
            pairs = alt[:3]
    return pairs


def write_reading_vii_excel(path: Path, section_vii_markdown: str, *, sheet_title: str) -> None:
    """
    Xuất mục VII sang Excel: questionName | answerName (đúng form import).
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Câu hỏi"
    hdrs = ("questionName", "answerName")
    for col, h in enumerate(hdrs, start=1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = Font(bold=True)
    rows = _parse_vii_excel_qa_rows(section_vii_markdown)
    if rows:
        for i, (q_raw, a_raw) in enumerate(rows, start=1):
            q = strip_markdown_for_excel_cell(q_raw)
            a = strip_markdown_for_excel_cell(a_raw)
            cell_q = ws.cell(row=1 + i, column=1, value=q)
            cell_a = ws.cell(row=1 + i, column=2, value=a)
            cell_q.alignment = Alignment(wrap_text=True, vertical="top")
            cell_a.alignment = Alignment(wrap_text=True, vertical="top")
    else:
        ws.merge_cells("A2:B2")
        c = ws["A2"]
        raw = (section_vii_markdown[:50000] if section_vii_markdown else "") or "(Rỗng)"
        c.value = strip_markdown_for_excel_cell(raw) if raw != "(Rỗng)" else raw
        c.alignment = Alignment(wrap_text=True, vertical="top")
    ws.column_dimensions["A"].width = 55
    ws.column_dimensions["B"].width = 72
    wb.save(str(path))


SYSTEM_READING = """Bạn là giảng viên lập trình, biên soạn bài đọc tiếng Việt theo mô hình **Storytelling in Tech**.

Luồng bắt buộc (nối mạch, không bỏ sót, không đảo thứ tự):
Đặt vấn đề → Phân tích → Giải pháp (đủ khái niệm, bản chất, khi nào dùng, cú pháp, giải thích cú pháp, lưu ý, đúc rút chuyên môn) → Ví dụ minh họa (code sạch, comment, cùng domain) → Giải quyết vấn đề ban đầu (quay lại phần I, code + phân tích) → Tổng kết (tóm tắt 3 ý, lỗi thường gặp, lưu ý sống còn, khái niệm hàn lâm ngắn, kết luận học thuật) → Câu hỏi kiểm tra (đúng 3 câu: Thông hiểu, Vận dụng có đoạn code, Phân tích mở rộng) → Tài liệu tham khảo.

Độ dài và trọng tâm (bắt buộc):
• **File DOCX (phần in):** **I → VI**, rồi **VII** (nội dung tài liệu tham khảo — trong bản Markdown gốc là `## VIII.`, tool đổi thành `## VII.` khi xuất Word cho đúng thứ tự). Mục câu hỏi (VII trong bản nháp) **không** in, xuất Excel. Khối DOCX mục tiêu **khoảng 4–5 trang A4** (Arial 12; **giãn dòng ~1,5** do tool đặt khi xuất Word; có hình minh họa nếu tool chèn). **Rút gọn tối đa:** chỉ giữ cốt lõi, không dông; tránh dưới **~3,5 trang** (thiếu bài) hoặc trên **~6 trang** (quá dài so mục tiêu).
• **Mục I:** khung 3 vấn đề; mỗi vấn đề **1–2 câu**; đoạn mở + hậu quả + câu hỏi trung tâm: cả mục I khoảng **6–10 câu**.
• **Mục II:** bóc điều kiện / ưu tiên; **không kể lại** câu chuyện mục I; **4–8 câu** hoặc **một** bảng nhỏ + **2–4 câu** nối.
• **Mục III:** ### **3–4 khối** (gộp ý khi được); mỗi ### **2–4 câu**; **một** khối ```python (**≤14 dòng**).
• **Mục IV:** ví dụ gọn + **một** khối ```python (**≤16 dòng**).
• **Mục V:** **một** khối ```python (**≤22 dòng**) + phân tích **3–5 câu**.
• **Mục VI:** **3** gạch tóm (mỗi gạch **1 câu** hoặc gọn); lỗi **3 gạch** (mỗi gạch **1 câu**); kết + thuật ngữ **1–3 câu**.
• **Mục VII:** (xuất Excel) 3 câu đủ chất; vận dụng **code ≤12 dòng**.
• **Mục VIII:** **3–4** nguồn. **Bắt buộc** mỗi dòng tham khảo là Markdown **liên kết**: `[mô tả ngắn tiếng Việt](https://URL-đầy-đủ)` (có thể thêm `- ` đầu dòng). **Cấm** chỉ dán URL trần không có phần mô tả trong ngoặc vuông — Word sẽ hiển thị **chữ có thể bấm mở link**.

Quy chuẩn:
• Một domain duy nhất xuyên suốt; không nhảy domain giữa các ví dụ.
• Không mở đầu kiểu “Hôm nay chúng ta học về…”. Tạo cảm giác vấn đề thực.
• **Chống lan man:** mỗi đoạn chỉ 1 ý chính; ưu tiên gạch đầu dòng ngắn; tránh kể thêm tình tiết phụ hoặc ví dụ “cho vui”. Nếu có chỗ dễ sa đà, hãy rút về: định nghĩa → khi dùng → dấu hiệu nhận biết → một ví dụ.
• Phần đầu gần gũi; phần tổng kết mới đậm học thuật (control flow, branching, boolean expression, business rule, mutual exclusivity, priority order, fallback case, maintainability — nêu ngắn gọn khi phù hợp).
• **Mã trong fence (Python, v.v.):** mọi **tên biến, tên hàm, khóa dict/dataclass, tên tham số** phải **tiếng Anh** (snake_case rõ nghĩa: `order_id`, `status`, `total_revenue`, `line_total`, `approved_count`). **Cấm** tên định danh romanized tiếng Việt hoặc không-ASCII trong code (ví dụ **không** dùng `ma`, `trang_thai`, `tong_doanh_thu`, `gia_tri`, `so_don_xu_ly`, `cho_xac_nhan` làm tên biến hay literal so sánh trong `if` — dùng `order_code`, `status`, `"pending_confirmation"`, …). **Comment** trong code và **nội dung chuỗi** trong `print(...)` / thông báo có thể tiếng Việt để khớp câu chuyện.
• Nếu trong yêu cầu có **phụ đề / nội dung đã trích từ video tham khảo**, bạn PHẢI tận dụng làm nguồn đồng hành với môn–session–lesson: bổ sung ví dụ, thuật ngữ, trọng tâm ôn tập; không bịa mâu thuẫn với phụ đề. Có thể ghi nguồn video ở mục VIII.
• **Thuật ngữ then chốt:** trong mọi đoạn văn (không áp dụng trong khối code), bọc các **từ khóa / khái niệm / cú pháp / mô hình** sinh viên phải nhớ bằng Markdown `**như thế này**` — khoảng vài cụm mỗi đoạn, ngắn gọn; không bôi đậm cả câu dài.
• **Biến/placeholder trong văn bản thường:** tránh biến 1 ký tự như `n`, `i` nếu không thật cần (dễ mơ hồ). Ưu tiên tên tiếng Anh rõ nghĩa: `num_rows`, `row_count`, `rows_len`, `max_index`, `end_index`… Khi cần nhắc biến / biểu thức ngắn trong câu (ví dụ “row chạy từ 0 đến n-1…”), hãy viết dạng `{{row_index}} chạy từ 0 đến {{row_count-1}}` (KHÔNG dùng backtick). Tool sẽ hiển thị {{…}} với nền xanh nhạt để đánh dấu biến/biểu thức.
• **Nhân vật & ảnh minh họa:** Tool có **2 nhân vật ảo cố định** (avatar) để ảnh mục I nhất quán: **nam** và **nữ** (đều ~đầu 20, bối cảnh văn phòng). Khi mục I có **một nhân vật người trung tâm** phù hợp (nhân viên văn phòng, dev, analyst, intern…), hãy chọn avatar **khớp giới** của nhân vật chính:  
  - Nam: **mái nấm ngố che trán (không rẽ ngôi)**, layer, xoăn/gợn lơi và rối nhẹ, **không undercut**; **áo sơ mi trắng cài cúc (không phanh hở ngực), có thể chỉ mở tối đa 1–2 cúc cổ**, không sơ vin, hơi oversized / gen Z hơi phố, quần jean xanh nhạt / wash trắng-xanh **không rách gối**; **ảnh mục I = 3D CGI điện ảnh siêu thực** (cinematic, bokeh) như avatar.  
  - Nữ: **váy trắng hở vai / cổ ống, thân váy nhiều lớp bèo/voan (tiered ruffle)**, trang sức mảnh; **cùng phong cách ảnh CGI** như trên.  
  Tên trong chuyện đặt tự do nhưng **giới + vai trò + trang phục** nên **gần** mô tả để ảnh mục I khớp. Nếu bài **không** có người trọng tâm hoặc dàn diễn đông, **không ép** avatar — ảnh I dùng biểu tượng/silhouette/tay, không mặt cụ thể.

Đầu ra: TOÀN BỘ bài bằng **Markdown** (tiếng Việt), không lời thoại meta, không emoji, không chỉ định màu trong văn bản; tool xuất DOCX sẽ **in đậm** các cụm đã bọc **…** trong Markdown.

Tiêu đề bài (đúng một dòng bắt đầu bằng #) PHẢI ĐÚNG mẫu sau (CHỮ HOA; hai cụm trong ngoặc kép tương ứng hai ẩn dụ song song; phần sau dấu hai chấm là phụ đề kiến thức):
# BÀI ĐỌC: "ẨN DỤ HOẶC MỆNH ĐỀ 1" VÀ "ẨN DỤ HOẶC MỆNH ĐỀ 2": PHẦN PHỤ ĐỀ MÔ TẢ TRỌNG TÂM (một dòng duy nhất, không xuống dòng)
Ví dụ chuẩn:
# BÀI ĐỌC: "NÚT PHANH KHẨN CẤP" VÀ "BƯỚC NHẢY THÔNG MINH": ĐIỀU KHIỂN LUỒNG DỮ LIỆU VỚI BREAK VÀ CONTINUE
Bám sát **môn, session, lesson** (và phụ đề video nếu có) để chọn hai cụm trong ngoặc kép và phần sau dấu hai chấm cho tiêu đề #.

Trình bày: dùng danh sách gạch đầu dòng (- ) cho liệt kê; khi so sánh hai khái niệm hãy dùng **bảng Markdown** (| cột 1 | cột 2 | và hàng phân cách |---|---|); có thể có đoạn so sánh 2 cột (vd. Break vs Continue).

Khung **I. ĐẶT VẤN ĐỀ** (bắt buộc — kiểu thực tế; **gọn** để tổng bài ~4–5 trang):
• Ngay dưới «## I. ĐẶT VẤN ĐỀ», mở bằng **1. Vấn đề: …** (câu mở IN HOA).
• Đoạn nhân vật + bối cảnh **vừa đủ**, không kể dài: họ tên Việt, vai trò, tổ chức, sự cố, nhiệm vụ (bỏ chi tiết phụ nếu không cần).
• **Ba khối** vấn đề — mỗi khối **1–2 câu**; **không** code ở mục I.
• «Nếu xử lý sai…» — **ba ý**, mỗi ý **1 câu** (tối đa 2 câu nếu thật cần).
• **Câu hỏi trung tâm:** **1–2 câu** có dấu hỏi, ẩn dụ khớp tiêu đề.

**Mục II → VI:** Luôn **nối mạch** cùng nhân vật, tổ chức và domain đã chọn ở mục I (không đổi sang ví dụ trừu tượng khác trừ khi lesson bắt buộc ví dụ khác — nếu vậy thì giải thích chuyển cảnh ngắn). II bóc tách ba tình huống thành điều kiện/ưu tiên; V phải **quay lại đúng** bài toán mở đầu và code xử lý.

Cấu trúc mục bắt buộc — mỗi mục chính là một dòng ## đúng dạng «## I.», «## II.», … «## VIII.» (số La Mã + dấu chấm ngay sau số, rồi tên mục; không đổi sang 1. hay chữ «Mục 1»). Ngay sau dòng tiêu đề # là:
## I. ĐẶT VẤN ĐỀ
## II. PHÂN TÍCH VẤN ĐỀ
## III. GIỚI THIỆU GIẢI PHÁP
(các tiểu mục ### **3–4 khối** — gộp tiểu mục khi trùng ý; **mỗi ### rất ngắn** theo giới hạn ở trên để tổng DOCX ~4–5 trang)
## IV. VÍ DỤ MINH HỌA
## V. GIẢI QUYẾT VẤN ĐỀ BAN ĐẦU
## VI. TỔNG KẾT VÀ LƯU Ý
## VII. BỘ CÂU HỎI KIỂM TRA
(đúng **3 mức** theo thứ tự: **Thông hiểu** → **Vận dụng** → **Phân tích**; mỗi mức **một khối** có đúng hai dòng nhãn để xuất Excel — viết nguyên văn **`**Câu hỏi:**`** rồi nội dung câu hỏi, xuống dòng **`**Đáp án:**`** rồi đáp án mẫu/gợi ý chấm đủ ý; **không** thêm dòng `---` hay `***` phân cách Markdown ngay sau đáp án; câu vận dụng phải có đoạn code trong khối ```python trong phần câu hỏi hoặc đáp án).
Ví dụ khung (bắt chước đúng nhãn **Câu hỏi:** / **Đáp án:**):
**Thông hiểu**

**Câu hỏi:** …

**Đáp án:** …

**Vận dụng**

**Câu hỏi:** …

**Đáp án:** …

**Phân tích**

**Câu hỏi:** …

**Đáp án:** …
**Lưu:** Mục «VII. BỘ CÂU HỎI…» **không** in DOCX — xuất **Excel**. Trong Markdown bạn vẫn viết mục tham khảo là **## VIII. TÀI LIỆU THAM KHẢO**; khi in DOCX tool **đánh lại số** thành **## VII.** để thứ tự La Mã trên giấy là I…VI→VII (tài liệu).
## VIII. TÀI LIỆU THAM KHẢO

Ví dụ đúng (từng dòng hoặc gạch đầu dòng):
- [Tài liệu Python về vòng lặp](https://docs.python.org/3/tutorial/controlflow.html)
- [Bài viết ngắn về best practice](https://example.com/bai-viet)

Dòng đầu của file Markdown sau phần tiêu đề # ở trên vẫn là ## I. ĐẶT VẤN ĐỀ (không lặp lại tiêu đề dạng #).

Mọi đoạn code Python đặt trong fence ```python (ghi rõ python) để DOCX tô màu cú pháp; ngôn ngữ khác vẫn hiển thị nền tối, chữ xám.
Code Python trong fence phải **đúng chỗ cách** (sau `for`, `in`, dấu phẩy, toán tử — ví dụ `for x in items:` chứ không viết dính `forx`).
Áp dụng **cùng quy tắc tên tiếng Anh** cho mọi khối ```python (kể cả mục VII — câu hỏi/đáp án có code).
Không dùng hình ảnh markdown ![ ]; không chèn URL ảnh."""


def _illustration_prompts(
    domain: str,
    core_concept: str,
    excerpt: str,
    *,
    model: str,
) -> tuple[str, str, str]:
    male_bible = READING_VIRTUAL_PROTAGONIST_MALE_EN
    female_bible = READING_VIRTUAL_PROTAGONIST_FEMALE_EN
    sys = (
        "You output EXACTLY 3 lines. Each line is ONE English prompt for an AI image model (detailed). "
        "Line 2–3: clean educational infographic / diagram look, cinematic lighting when it helps; "
        "Line 1: hyper-realistic cinematic 3D CGI per CHARACTER BIBLE (male or female—same render quality, outfit locked in bible). "
        "Vietnam-relevant professional setting when the text implies it."
        "\nCRITICAL — on-screen text: Image generators mangle Vietnamese spelling. "
        "Do NOT put Vietnamese words on monitors, signs, or labels. "
        "Prefer NO readable text: icons, symbols, colors, arrows, flowchart shapes, split panels. "
        "If labels unavoidable: ONLY short ENGLISH words (STOP, SKIP, IF, CASE 1–3, LIMIT). "
        "No emoji. No watermark."
        "\nCHARACTER BIBLES (Line 1 only): Choose ONE fixed virtual avatar if excerpt clearly centers ONE human workplace "
        "protagonist with a clear gender and office context; otherwise use abstract icons/silhouettes/hands-only.\n"
        f"- MALE AVATAR (use only if lead is male): {male_bible}\n"
        f"- FEMALE AVATAR (use only if lead is female): {female_bible}"
        "\nLine 1 — Section I (problem / crisis): When a CHARACTER BIBLE applies, show THAT SAME protagonist as in the bible—"
        "hyper-realistic cinematic 3D CGI, locked outfit, identical lighting/render language for male or female lead. "
        "Subject at workstation in crisis; urgent monitors; "
        "THREE-PART metaphor (threat stop / skip noise / limit exit) matching story. "
        "If Bible does NOT apply, no detailed face — icons, hands, or abstract figures only."
        "\nLine 2 — Section II (analysis / breakdown): Abstract diagram — decision tree, three parallel tracks, "
        "priority arrows, or condition matrix mapping story situations to technical choices; same domain mood; "
        "no Vietnamese text; no recurring face."
        "\nLine 3 — Section III (solution / mechanics): Clean flowchart — loops, diamonds, branch / break / continue metaphor; "
        "same palette; no Vietnamese on image; no recurring face."
        "\nNo numbering prefixes. No extra commentary."
    )
    user = (
        f"Domain / môn: {domain}\nLesson focus: {core_concept}\n\n"
        f"Excerpt (mục I → II nếu có — nhân vật & tình huống):\n{excerpt[:5000]}\n\n"
        "If there is a single clear lead: pick the matching avatar (MALE or FEMALE) and Line 1 MUST embed those visual traits. "
        "If unclear/no lead/ensemble: do NOT pick an avatar and do NOT show a detailed face."
    )
    raw, _ = complete_chat(
        [ChatMessage(role="system", content=sys), ChatMessage(role="user", content=user)],
        model=model,
        temperature=0.38,
        max_tokens=900,
    )
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip() and not ln.strip().startswith("#")]
    if len(lines) >= 3:
        return lines[0][:1200], lines[1][:1200], lines[2][:1200]
    if len(lines) == 2:
        mid = lines[1][:1200]
        return lines[0][:1200], mid, mid
    parts = [p.strip() for p in raw.replace("\n", ";").split(";") if len(p.strip()) > 22]
    if len(parts) >= 3:
        return parts[0][:1200], parts[1][:1200], parts[2][:1200]
    fb_a = (
        f"Hyper-realistic cinematic 3D CGI: {READING_VIRTUAL_PROTAGONIST_MALE_EN} At operations desk, urgent monitors; "
        f"three-part conveyor metaphor STOP SKIP LIMIT icons; domain {domain}, lesson {core_concept}"
    )
    fb_b = (
        f"Abstract analysis infographic: three branching paths or decision matrix, priority arrows, "
        f"condition icons, no Vietnamese; lesson {core_concept}"
    )
    fb_c = (
        f"Technical flowchart: loop control, diamonds, branches; English micro-labels or no text; {core_concept}"
    )
    return fb_a, fb_b, fb_c


def _strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", s)


def _split_inline_md_bold_chunks(text: str) -> list[tuple[bool, str]]:
    """Tách chuỗi theo cặp **…** → [(is_key_term, segment), …]."""
    if not text:
        return []
    parts: list[tuple[bool, str]] = []
    idx = 0
    n = len(text)
    while idx < n:
        j = text.find("**", idx)
        if j < 0:
            parts.append((False, text[idx:]))
            break
        if j > idx:
            parts.append((False, text[idx:j]))
        k = text.find("**", j + 2)
        if k < 0:
            parts.append((False, text[j:]))
            break
        inner = text[j + 2 : k]
        if inner:
            parts.append((True, inner))
        idx = k + 2
    return parts


_MD_INLINE_LINK = re.compile(r"\[([^\]]*)\]\(([^)]*)\)")
# Marker for variables / tiny expressions in prose, to render as "chips" in DOCX.
# Examples: {{n}}, {{row}}, {{n-1}}, {{i+1}} (we keep it conservative: var +/- integer only).
_MD_INLINE_VAR = re.compile(r"\{\{\s*([A-Za-z_][A-Za-z0-9_]*\s*(?:[+-]\s*\d+)?)\s*\}\}")


def _iter_inline_md_segments(
    text: str,
) -> list[tuple[str, str] | tuple[str, str, str]]:
    """
    Tách đoạn theo [nhãn](url) rồi **…** trong phần còn lại.
    Phần tử: ('plain', s) | ('bold', s) | ('link', nhãn, url).
    """
    out: list[tuple[str, str] | tuple[str, str, str]] = []
    pos = 0
    for m in _MD_INLINE_LINK.finditer(text):
        if m.start() > pos:
            for is_key, seg in _split_inline_md_bold_chunks(text[pos : m.start()]):
                if seg:
                    out.append(("bold", seg) if is_key else ("plain", seg))
        out.append(("link", m.group(1), m.group(2)))
        pos = m.end()
    if pos < len(text):
        for is_key, seg in _split_inline_md_bold_chunks(text[pos:]):
            if seg:
                out.append(("bold", seg) if is_key else ("plain", seg))
    return out


def _hyperlink_run_rpr(*, size_pt: float, bold: bool) -> OxmlElement:
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), READING_DOC_FONT)
    r_fonts.set(qn("w:hAnsi"), READING_DOC_FONT)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(round(float(size_pt) * 2))))
    r_pr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    return r_pr


def _append_hyperlink(paragraph, display_text: str, url: str, *, size_pt: float) -> None:
    """Chèn hyperlink Word (chữ xanh, gạch chân); lỗi URI → chữ thường."""
    url = (url or "").strip()
    display = (display_text or "").strip()
    if not url:
        if display:
            r = paragraph.add_run(display)
            _style_body_run(r, size_pt=size_pt)
        return
    if not display:
        display = url
    try:
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    except Exception:
        r = paragraph.add_run(display)
        _style_body_run(r, size_pt=size_pt)
        return
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    for is_key, seg in _split_inline_md_bold_chunks(display):
        if not seg:
            continue
        new_run = OxmlElement("w:r")
        new_run.append(_hyperlink_run_rpr(size_pt=size_pt, bold=is_key))
        t_el = OxmlElement("w:t")
        if seg.startswith(" ") or seg.endswith(" "):
            t_el.set(qn("xml:space"), "preserve")
        t_el.text = seg
        new_run.append(t_el)
        hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_run_shading(run, fill_hex: str) -> None:
    """Set background shading for a run (Word)."""
    try:
        r_el = run._r  # type: ignore[attr-defined]
        r_pr = r_el.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        r_pr.append(shd)
    except Exception:
        return


_VAR_CHIP_FILL = "DCEBFF"  # light blue


def _append_text_with_var_chips(p, s: str, *, size_pt: float, bold: bool) -> None:
    """
    Render plain text with {{var}} markers as 'chips' (light-blue background),
    without printing the braces.
    """
    if not s:
        return
    pos = 0
    for m in _MD_INLINE_VAR.finditer(s):
        if m.start() > pos:
            r = p.add_run(s[pos : m.start()])
            _style_body_run(r, size_pt=size_pt, bold=bold)
        var = (m.group(1) or "").strip()
        if var:
            r2 = p.add_run(var)
            _style_body_run(r2, size_pt=size_pt, bold=bold)
            _set_run_shading(r2, _VAR_CHIP_FILL)
        pos = m.end()
    if pos < len(s):
        r3 = p.add_run(s[pos:])
        _style_body_run(r3, size_pt=size_pt, bold=bold)


def _h2_roman_section_key(line: str) -> str | None:
    """Lấy nhãn La Mã từ tiêu đề cấp 2 kiểu «## I. ĐẶT VẤN ĐỀ» → «I»."""
    m = re.match(r"^##\s+([IVXLCDM]+)\.", line.strip(), re.I)
    if not m:
        return None
    return m.group(1).upper()


_BLACK = RGBColor(0x00, 0x00, 0x00)

# Font văn bản chính (đề, đoạn, bảng, heading do tool gán run — không áp dụng khối code).
READING_DOC_FONT = "Arial"

# Khối code: nền tối kiểu IDE; màu chữ gần theme tối (comment xanh, từ khóa hồng, hàm tím, v.v.)
_CODE_BLOCK_FILL = "1E1E1E"
_CLR_CODE_COMMENT = RGBColor(0x9C, 0xDC, 0xFE)  # comment — xanh nhạt
_CLR_CODE_KEYWORD = RGBColor(0xD1, 0x6D, 0x9A)  # for, if, break — hồng/đỏ nhạt
_CLR_CODE_FUNCTION = RGBColor(0xC5, 0x86, 0xC0)  # gọi hàm — tím
_CLR_CODE_STRING = RGBColor(0xCE, 0x91, 0x78)  # chuỗi — cam
_CLR_CODE_NUMBER = RGBColor(0xB5, 0xCE, 0xA8)  # số — xanh lá nhạt
_CLR_CODE_DEFAULT = RGBColor(0xD4, 0xD4, 0xD4)  # biến, toán tử

_PY_RESERVED: frozenset[str] = frozenset(keyword.kwlist) | frozenset({"True", "False", "None"})


def _configure_document_defaults(doc: Document) -> None:
    """Lề A4 kiểu văn bản Word thông dụng (VN); font mặc định Arial 12pt."""
    sec = doc.sections[0]
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = READING_DOC_FONT
    normal.font.size = Pt(12)


def _style_body_run(run, *, size_pt: float = 12.0, bold: bool = False) -> None:
    run.font.color.rgb = _BLACK
    run.font.name = READING_DOC_FONT
    run.font.size = Pt(int(size_pt))
    run.bold = bold


def _style_key_term_run(run, *, size_pt: float = 12.0) -> None:
    """Từ khóa: in đậm (không tô nền)."""
    run.font.color.rgb = _BLACK
    run.font.name = READING_DOC_FONT
    run.font.size = Pt(int(size_pt))
    run.bold = True


def _set_paragraph_shading(p, fill_hex: str) -> None:
    p_pr = p._element.get_or_add_pPr()
    for shd in list(p_pr.findall(qn("w:shd"))):
        p_pr.remove(shd)
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), fill_hex)
    el.set(qn("w:val"), "clear")
    p_pr.append(el)


def _paragraph_has_code_block_shading(paragraph) -> bool:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return False
    for shd in p_pr.findall(qn("w:shd")):
        fill = shd.get(qn("w:fill"))
        if fill and fill.upper() == _CODE_BLOCK_FILL.upper():
            return True
    return False


def _style_code_run_on_dark(run, *, color: RGBColor = _CLR_CODE_DEFAULT) -> None:
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = color


def _token_next_non_comment_newline(tokens: list, i: int):
    j = i + 1
    while j < len(tokens):
        t = tokens[j]
        if t.type in (tokenize.NL, tokenize.NEWLINE, tokenize.COMMENT):
            j += 1
            continue
        return t
    return None


def _char_index_in_code(code: str, line_1based: int, col_0based: int) -> int:
    """Chuyển (dòng 1-based, cột 0-based) của `tokenize` thành offset byte trong `code`."""
    if line_1based < 1:
        return 0
    lines = code.splitlines(keepends=True)
    idx = 0
    for li in range(line_1based - 1):
        if li < len(lines):
            idx += len(lines[li])
    idx += col_0based
    return min(max(0, idx), len(code))


def _python_token_color(tok: tokenize.TokenInfo, tokens: list, i: int) -> RGBColor:
    if tok.type == tokenize.NAME:
        if tok.string in _PY_RESERVED:
            return _CLR_CODE_KEYWORD
        nxt = _token_next_non_comment_newline(tokens, i)
        if nxt and nxt.type == tokenize.OP and nxt.string == "(":
            return _CLR_CODE_FUNCTION
        return _CLR_CODE_DEFAULT
    if tok.type == tokenize.COMMENT:
        return _CLR_CODE_COMMENT
    if tok.type == tokenize.STRING:
        return _CLR_CODE_STRING
    if tok.type == tokenize.NUMBER:
        return _CLR_CODE_NUMBER
    return _CLR_CODE_DEFAULT


def _add_python_syntax_runs(p, code: str) -> None:
    """Tô màu Python; luôn chèn phần nguồn giữa các token (khoảng trắng, xuống dòng) vì `tokenize` không tạo token cho space."""
    try:
        tokens = list(tokenize.generate_tokens(io.StringIO(code).readline))
    except tokenize.TokenError:
        r = p.add_run(code)
        _style_code_run_on_dark(r)
        return
    cursor = 0
    for i, tok in enumerate(tokens):
        idx_s = _char_index_in_code(code, tok.start[0], tok.start[1])
        idx_e = _char_index_in_code(code, tok.end[0], tok.end[1])
        if cursor < idx_s:
            gap = code[cursor:idx_s]
            r = p.add_run(gap)
            _style_code_run_on_dark(r, color=_CLR_CODE_DEFAULT)
        if tok.type == tokenize.ENDMARKER:
            cursor = idx_e
            break
        if tok.type == tokenize.ENCODING:
            cursor = idx_e
            continue
        if idx_s >= idx_e:
            cursor = idx_e
            continue
        piece = code[idx_s:idx_e]
        color = _python_token_color(tok, tokens, i)
        r = p.add_run(piece)
        _style_code_run_on_dark(r, color=color)
        cursor = idx_e
    if cursor < len(code):
        r = p.add_run(code[cursor:])
        _style_code_run_on_dark(r, color=_CLR_CODE_DEFAULT)


def _add_code_block_paragraph(doc: Document, code: str, fence_lang: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_shading(p, _CODE_BLOCK_FILL)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.left_indent = Inches(0.14)
    pf.right_indent = Inches(0.14)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    lang = (fence_lang or "").strip().lower()
    if lang in ("python", "py", "python3", ""):
        _add_python_syntax_runs(p, code)
    else:
        r = p.add_run(code)
        _style_code_run_on_dark(r)


def _body_paragraph_format(p, *, list_indent: bool = False) -> None:
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if list_indent:
        pf.left_indent = Inches(0.28)
        pf.first_line_indent = Inches(-0.18)


def _heading_paragraph_format(p, *, level: int) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(16 if level <= 2 else 12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    pf.widow_control = True


def _populate_inline_md_runs(
    p,
    text: str,
    *,
    size_pt: float = 12.0,
    heading_bold_non_key: bool = False,
    prefix: str | None = None,
) -> None:
    """Thêm run vào đoạn; **…** → in đậm; [nhãn](url) → hyperlink Word."""
    if prefix:
        r0 = p.add_run(prefix)
        _style_body_run(r0, size_pt=size_pt, bold=heading_bold_non_key)
    for item in _iter_inline_md_segments(text):
        if item[0] == "plain":
            _append_text_with_var_chips(p, item[1], size_pt=size_pt, bold=heading_bold_non_key)
        elif item[0] == "bold":
            seg = item[1]
            pos = 0
            for m in _MD_INLINE_VAR.finditer(seg):
                if m.start() > pos:
                    r = p.add_run(seg[pos : m.start()])
                    _style_key_term_run(r, size_pt=size_pt)
                var = (m.group(1) or "").strip()
                if var:
                    r2 = p.add_run(var)
                    _style_key_term_run(r2, size_pt=size_pt)
                    _set_run_shading(r2, _VAR_CHIP_FILL)
                pos = m.end()
            if pos < len(seg):
                r3 = p.add_run(seg[pos:])
                _style_key_term_run(r3, size_pt=size_pt)
        else:
            _, label, url = item
            _append_hyperlink(p, label, url, size_pt=size_pt)


def _add_styled_heading(
    doc: Document,
    text: str,
    level: int,
    *,
    size_pt: int,
    center: bool = False,
) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _populate_inline_md_runs(p, text, size_pt=float(size_pt), heading_bold_non_key=True)
    _heading_paragraph_format(p, level=level)


def _is_table_row_line(s: str) -> bool:
    s = s.strip()
    return len(s) >= 2 and s.startswith("|") and s.endswith("|")


def _table_cells_from_line(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|") and s.endswith("|"):
        s = s[1:-1]
    return [c.strip() for c in s.split("|")]


def _is_md_table_separator(cells: list[str]) -> bool:
    """Nhận dạng hàng |---|---| trong bảng Markdown."""
    if not cells:
        return False
    for c in cells:
        t = "".join(ch for ch in c if ch not in " \t")
        if not t or any(ch not in "-:" for ch in t) or "-" not in t:
            return False
    return True


def _parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]] | None, int]:
    if start >= len(lines) or not _is_table_row_line(lines[start]):
        return None, start
    rows: list[list[str]] = []
    j = start
    while j < len(lines) and _is_table_row_line(lines[j]):
        cells = _table_cells_from_line(lines[j])
        if _is_md_table_separator(cells):
            j += 1
            continue
        rows.append(cells)
        j += 1
    return rows if rows else None, j


def _add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    norm = [r + [""] * (n_cols - len(r)) if len(r) < n_cols else r[:n_cols] for r in rows]
    tbl = doc.add_table(rows=len(norm), cols=n_cols)
    tbl.style = "Table Grid"
    for ri, row_cells in enumerate(norm):
        is_header = ri == 0
        for ci, cell_text in enumerate(row_cells):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _populate_inline_md_runs(
                p,
                cell_text,
                size_pt=11.0,
                heading_bold_non_key=is_header,
            )
            _body_paragraph_format(p)


def _add_centered_picture_stream(doc: Document, stream: io.BytesIO, *, width_in: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(stream, width=Inches(width_in))


def _add_picture_robust(doc: Document, blob: bytes, *, width_in: float = 5.8) -> bool:
    """Chèn ảnh căn giữa trang; WebP/RGBA chuyển PNG qua Pillow nếu python-docx không đọc được."""
    try:
        _add_centered_picture_stream(doc, io.BytesIO(blob), width_in=width_in)
        return True
    except Exception:
        pass
    try:
        from PIL import Image

        im = Image.open(io.BytesIO(blob))
        if im.mode == "P":
            im = im.convert("RGBA")
        if im.mode == "RGBA":
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=im.split()[3])
            im = bg
        elif im.mode == "LA":
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=im.split()[1])
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")
        buf = io.BytesIO()
        im.save(buf, format="PNG")
        buf.seek(0)
        _add_centered_picture_stream(doc, buf, width_in=width_in)
        return True
    except Exception:
        return False


def _force_all_text_black(doc: Document) -> None:
    for p in doc.paragraphs:
        if _paragraph_has_code_block_shading(p):
            continue
        for r in p.runs:
            r.font.color.rgb = _BLACK
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = _BLACK


def markdown_to_docx(
    md: str,
    *,
    section_images: dict[str, bytes] | None = None,
    orphan_images: list[bytes] | None = None,
) -> Document:
    doc = Document()
    _configure_document_defaults(doc)
    lines = md.splitlines()
    i = 0
    code_mode = False
    code_buf: list[str] = []
    para_buf: list[str] = []
    reading_main_title_placed = False
    inserted_section_keys: set[str] = set()
    code_fence_lang = ""

    def flush_para() -> None:
        text = "\n".join(para_buf).strip()
        para_buf.clear()
        if not text:
            return
        text = _strip_trailing_horizontal_rules(text)
        if not text:
            return
        p = doc.add_paragraph()
        _populate_inline_md_runs(p, text, size_pt=12.0)
        _body_paragraph_format(p)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_para()
            if not code_mode:
                code_mode = True
                code_buf = []
                rest = stripped[3:].strip()
                code_fence_lang = (rest.split()[0].lower() if rest else "") or ""
            else:
                _add_code_block_paragraph(doc, "\n".join(code_buf), code_fence_lang)
                code_mode = False
                code_buf = []
                code_fence_lang = ""
            i += 1
            continue
        if code_mode:
            code_buf.append(line)
            i += 1
            continue

        tbl, ni = _parse_md_table(lines, i)
        if tbl is not None:
            flush_para()
            _add_markdown_table(doc, tbl)
            gap = doc.add_paragraph()
            _body_paragraph_format(gap)
            i = ni
            continue

        num_m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if num_m:
            flush_para()
            p = doc.add_paragraph()
            _populate_inline_md_runs(
                p,
                num_m.group(2),
                size_pt=12.0,
                prefix=f"{num_m.group(1)}. ",
            )
            _body_paragraph_format(p, list_indent=True)
            i += 1
            continue

        bull_m = re.match(r"^[-*•]\s+(.+)$", stripped)
        if bull_m:
            flush_para()
            p = doc.add_paragraph()
            _populate_inline_md_runs(p, bull_m.group(1), size_pt=12.0, prefix="• ")
            _body_paragraph_format(p, list_indent=True)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_para()
            _add_styled_heading(doc, stripped[4:].strip(), 3, size_pt=12)
            i += 1
            continue
        if stripped.startswith("## "):
            flush_para()
            _add_styled_heading(doc, stripped[3:].strip(), 2, size_pt=14)
            sec_key = _h2_roman_section_key(stripped)
            if (
                sec_key
                and section_images
                and sec_key in section_images
                and sec_key not in inserted_section_keys
            ):
                if not _add_picture_robust(doc, section_images[sec_key], width_in=5.5):
                    pf = doc.add_paragraph()
                    rf = pf.add_run(f"[Không chèn được minh họa mục {sec_key}.]")
                    _style_body_run(rf, size_pt=10.0)
                inserted_section_keys.add(sec_key)
            i += 1
            continue
        if stripped.startswith("# "):
            flush_para()
            raw_h1 = stripped[2:].strip()
            do_center = not reading_main_title_placed
            _add_styled_heading(doc, raw_h1, 1, size_pt=16, center=do_center)
            if do_center:
                reading_main_title_placed = True
            i += 1
            continue
        if not stripped:
            flush_para()
            i += 1
            continue
        if _HR_ONLY_LINE.match(stripped):
            flush_para()
            i += 1
            continue
        para_buf.append(line)
        i += 1
    flush_para()

    tail_blobs: list[bytes] = []
    if section_images:
        for sk, blob in section_images.items():
            if sk not in inserted_section_keys:
                tail_blobs.append(blob)
    if orphan_images:
        tail_blobs.extend(orphan_images)
    if tail_blobs:
        _add_styled_heading(doc, "Minh họa bổ sung (ảnh AI)", 2, size_pt=14)
        p = doc.add_paragraph()
        _populate_inline_md_runs(
            p,
            "Các ảnh chưa gắn được vào đúng mục La Mã trong file (khác định dạng tiêu đề), hoặc ảnh dự phòng.",
            size_pt=11.0,
        )
        _body_paragraph_format(p)
        for idx, blob in enumerate(tail_blobs, start=1):
            p2 = doc.add_paragraph()
            r2 = p2.add_run(f"Hình bổ sung {idx}.")
            _style_body_run(r2, size_pt=12.0)
            _body_paragraph_format(p2)
            if not _add_picture_robust(doc, blob, width_in=5.8):
                pe = doc.add_paragraph()
                re_ = pe.add_run(
                    f"[Không chèn được hình {idx} — kiểm tra Pillow hoặc dữ liệu ảnh từ API.]"
                )
                _style_body_run(re_, size_pt=11.0)

    _force_all_text_black(doc)
    return doc


@dataclass
class ReadingDocParams:
    """Đầu vào chính: môn, session, lesson; video YouTube (tuỳ chọn) — dùng phụ đề làm nguồn."""

    subject: str
    session: str
    lesson: str
    session_stt: str
    lesson_stt: str
    video_url: str | None
    learning_goals: str
    references_hint: str
    text_model: str
    image_model: str
    generate_illustrations: bool
    output_docx: Path
    output_xlsx: Path
    technology: str = ""
    audience: str = ""


def run_reading_generation(
    params: ReadingDocParams,
    *,
    on_progress: Callable[[str], None] | None = None,
) -> tuple[bool, str]:
    subject = (params.subject or "").strip()
    session = (params.session or "").strip()
    lesson = (params.lesson or "").strip()
    if not subject:
        return False, "Thiếu môn."
    if not session:
        return False, "Thiếu session."
    if not lesson:
        return False, "Thiếu lesson."
    if not params.output_docx:
        return False, "Thiếu đường dẫn file DOCX đầu ra."
    if not params.output_xlsx:
        return False, "Thiếu đường dẫn file Excel đầu ra."

    video_notes = ""
    video_url = (params.video_url or "").strip()
    if video_url:
        video_notes, transcript_error = fetch_youtube_transcript_plain(video_url)
        if transcript_error or not video_notes:
            return False, (
                transcript_error
                or "Không lấy được phụ đề video — bỏ link hoặc dùng video YouTube có phụ đề tiếng Việt / tiếng Anh."
            )

    text_model = resolve_model(params.text_model or None)
    goals = (params.learning_goals or "").strip() or DEFAULT_LEARNING_GOALS
    tech = (params.technology or "").strip()
    if not tech:
        tech = f"Tự chọn ngôn ngữ/công nghệ phù hợp môn «{subject}» và lesson «{lesson}» (nếu là lập trình thì chọn đúng stack thường dạy của môn)."
    audience = (params.audience or "").strip() or "Sinh viên"

    framing = f"Môn «{subject}» — Session «{session}» — Lesson «{lesson}»"

    user = f"""Thông tin khung bài đọc:

• Môn: {subject}
• Session: {session}
• Lesson (chủ đề / kiến thức trọng tâm của bài đọc): {lesson}
• Ngôn ngữ / công nghệ (cho ví dụ code): {tech}
• Đối tượng: {audience}
• Cốt lõi “một domain” xuyên suốt truyện: lấy từ môn + lesson — không nhảy sang lĩnh vực không liên quan.

Mục tiêu sau bài đọc (bám sát, có thể diễn đạt lại mạch lạc):
{goals}
"""
    if video_notes:
        user += f"""
---
PHỤ ĐỀ VIDEO THAM KHẢO (đã trích từ YouTube — bạn PHẢI khai thác, không mâu thuẫn nội dung này):
{video_notes}
---
Link gốc (đưa vào mục VIII nếu thích): {video_url}
"""

    if (params.references_hint or "").strip():
        user += "\nGợi ý tài liệu tham khảo thêm (mục VIII):\n"
        user += params.references_hint.strip()

    user += (
        "\n\nSoạn TOÀN BỘ bài theo SYSTEM — **phần in DOCX (I–VI, rồi mục tham khảo hiển thị là VII) khoảng 4–5 trang A4**, **rút gọn**, chỉ cốt lõi; mục VIII chỉ dùng dòng dạng `[mô tả](URL)`. "
        f"Code ví dụ thống nhất với: {tech}; **tên biến/hàm/identifier trong code chỉ tiếng Anh** (comment và chuỗi print có thể tiếng Việt).\n"
        "Phần VII: đúng 3 mức Thông hiểu / Vận dụng / Phân tích; mỗi mức phải có dòng **Câu hỏi:** và **Đáp án:** (để tool map cột questionName / answerName).\n"
        "Dòng tiêu đề # PHẢI theo mẫu BÀI ĐỌC: \"…\" VÀ \"…\": … như ví dụ trong SYSTEM; "
        f"hai ẩn dụ và phụ đề kiến thức phải bám sát: {framing}"
        + (" và phản ánh nội dung đã có trong phụ đề video." if video_notes else ".")
    )

    if on_progress:
        on_progress(
            "Bước 1/3: Đang gọi model soạn bài (DOCX ~4–5 trang, rút gọn; thường 2–10 phút)…"
        )

    md, _ = complete_chat(
        [
            ChatMessage(role="system", content=SYSTEM_READING),
            ChatMessage(role="user", content=user),
        ],
        model=text_model,
        temperature=0.38,
        max_tokens=20000,
        timeout_s=540.0,
    )

    if not md or len(md) < 400:
        return False, "Model trả nội dung quá ngắn hoặc rỗng."

    markdown_for_docx, section_vii_markdown = split_markdown_remove_vii_for_docx(md)
    markdown_for_docx = renumber_references_viii_to_vii_for_docx(markdown_for_docx)

    images: list[bytes] = []
    section_map: dict[str, bytes] = {}
    if params.generate_illustrations:
        primary_image_model = (params.image_model or "").strip() or "black-forest-labs/flux.2-pro"
        if on_progress:
            on_progress("Bước 2/3: Đang tạo prompt minh họa…")
        try:
            illustration_prompt_1, illustration_prompt_2, illustration_prompt_3 = (
                _illustration_prompts(subject, lesson, markdown_for_docx, model=text_model)
            )
        except Exception:
            illustration_prompt_1, illustration_prompt_2, illustration_prompt_3 = (
                f"Tech crisis infographic, icons only, domain {subject} {lesson}",
                f"Decision breakdown diagram three paths, lesson {lesson}, no Vietnamese text",
                f"Flowchart loop control {lesson}, English micro-labels or no text",
            )
        if on_progress:
            on_progress(
                "Bước 3/3: Đang tạo 3 ảnh minh họa song song (mục I, II, III — mỗi ảnh có thể 1–4 phút)…"
            )
        with ThreadPoolExecutor(max_workers=3) as pool:
            image_job_1 = pool.submit(
                _reading_first_image_blob, illustration_prompt_1, primary_image_model
            )
            image_job_2 = pool.submit(
                _reading_first_image_blob, illustration_prompt_2, primary_image_model
            )
            image_job_3 = pool.submit(
                _reading_first_image_blob, illustration_prompt_3, primary_image_model
            )
            illustration_image_1 = image_job_1.result()
            illustration_image_2 = image_job_2.result()
            illustration_image_3 = image_job_3.result()
            for key, blob in zip(
                ("I", "II", "III"),
                (illustration_image_1, illustration_image_2, illustration_image_3),
                strict=True,
            ):
                if blob:
                    section_map[key] = blob
            images = [
                b
                for b in (illustration_image_1, illustration_image_2, illustration_image_3)
                if b
            ]

    doc = markdown_to_docx(
        markdown_for_docx,
        section_images=section_map if section_map else None,
    )
    if params.generate_illustrations and not images:
        p = doc.add_paragraph()
        r = p.add_run(
            "[Ghi chú] Chưa nhận được ảnh từ API. Kiểm tra model ảnh (output_modalities có image trên OpenRouter) "
            "hoặc thử model khác; nội dung bài đọc vẫn đầy đủ. Ảnh không phải ảnh chụp code — đoạn code chỉ nên là văn bản."
        )
        _style_body_run(r, size_pt=12.0)
    _force_all_text_black(doc)
    output_stem = reading_output_stem(subject, params.session_stt, params.lesson_stt)
    excel_sheet_title = f"Bộ câu hỏi (VII) — {output_stem}"
    try:
        params.output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(params.output_docx))
    except Exception as e:
        return False, f"Không ghi được DOCX: {e}"
    try:
        write_reading_vii_excel(
            params.output_xlsx,
            section_vii_markdown or "",
            sheet_title=excel_sheet_title,
        )
    except Exception as e:
        return False, f"Đã lưu DOCX nhưng không ghi được Excel: {e}"

    return True, (
        f"DOCX:\n{params.output_docx.resolve()}\n\n"
        f"Excel (questionName | answerName):\n{params.output_xlsx.resolve()}"
    )
